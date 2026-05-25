"""
Case Study Generator — implements the Case Study Generator Prompts v2 design
document for the CBC-India AGK Case Study Suite.

Three case types:
    - Lesson-Drawing  (3,000-7,000 words)
    - Decision-Forcing (2,500-5,000 words)
    - Caselet         (under 2,000 words)

Public entry points are grouped by stage:
    Stage 1 (intake)             -> assemble_source_bundle()
    Stage 2 (source processing)  -> run_source_processing()
    Stage 3 (drafting)           -> draft_case_section() / draft_caselet()
    Stage 4 (compliance)         -> run_compliance_passes()
    Stage 5 (teaching note)      -> generate_teaching_note()
    Stage 6 (export)             -> build_case_docx() / build_case_pdf() /
                                    build_teaching_note_docx()

All AI calls go through utils.call_openai_api with temperature=0.1 and seed=42
for deterministic output, mirroring the writing assistant pattern.
"""

import json
import os
import re
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

from utils import call_openai_api, sanitize_text_for_pdf
from assessment_criteria import KCM_COMPETENCIES


# ---------------------------------------------------------------------------
# Master system prompt + case-type metadata
# ---------------------------------------------------------------------------

MASTER_SYSTEM_PROMPT = """You are a professional case study writer for the
Capacity Building Commission, Government of India, helping authors produce case
studies for the Amrit Gyaan Kosh knowledge repository.

Non-negotiable writing rules:
- Write in BRITISH ENGLISH only (organise, programme, centre, analyse, colour, etc.).
- Write in the PAST TENSE throughout. Direct quotes may be in any tense.
  General principles inside a Lessons Learnt section may use the present tense.
- Use plain, precise, factual prose. No flowery, emotional or promotional
  language. No superlatives.
- Do NOT describe any individual as a hero, visionary or exceptional person.
  Describe what people did and what resulted from it.
- Use adjectives sparingly. If removing the adjective does not change the
  factual meaning, leave it out.
- Avoid em-dashes. Default to commas, full stops or sentence restructuring.
- Stay politically neutral. Do not comment on parties, electoral outcomes or
  ideological positions.
- Do NOT fabricate facts, quotes, statistics, dates, names or outcomes. If the
  source material does not state it, do not include it. If something is
  uncertain, omit it or mark it explicitly as uncertain.
- Every direct quote must be verbatim from the source material and attributed
  to a named speaker.
"""

CASE_TYPES = {
    "lesson_drawing": {
        "label": "Lesson-Drawing Case Study",
        "word_min": 3000,
        "word_max": 7000,
        "default_target": 5000,
        "description": (
            "A detailed narrative account of how an initiative was designed, "
            "implemented and experienced. Traces events in chronological order, "
            "captures multiple stakeholder perspectives and draws transferable "
            "lessons at the end. Does not ask the reader to make a decision."
        ),
        "sections": [
            ("hook", "Hook"),
            ("introduction", "Introduction and Context Setting"),
            ("dilemma", "Development Challenge and Central Dilemma"),
            ("delivery_challenges", "Delivery Challenges"),
            ("implementation", "Chronological Implementation"),
            ("lessons", "Lessons Learnt"),
            ("closing", "Closing"),
        ],
    },
    "decision_forcing": {
        "label": "Decision-Forcing Case Study",
        "word_min": 2500,
        "word_max": 5000,
        "default_target": 3500,
        "description": (
            "Places the reader in the position of the protagonist at a critical "
            "moment of choice. Builds up context, pressures and options but "
            "stops before revealing what actually happened. The reader analyses "
            "the situation and arrives at their own recommendation."
        ),
        "sections": [
            ("executive_summary", "Executive Summary"),
            ("introduction", "Introduction"),
            ("background", "Background"),
            ("narrative", "Narrative"),
            ("dilemma", "Dilemma or Puzzle"),
            ("epilogue", "Epilogue (Optional)"),
            ("study_questions", "Study Questions"),
        ],
    },
    "caselet": {
        "label": "Caselet",
        "word_min": 600,
        "word_max": 2000,
        "default_target": 1500,
        "description": (
            "A short, focused account of a specific situation, challenge or "
            "moment. Introduces a context, describes what happened or what "
            "decision was faced, and closes with one or two discussion "
            "questions. Useful as a brief illustration."
        ),
        "sections": [
            ("caselet", "Full Caselet"),
        ],
    },
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _truncate(text, limit):
    if not text:
        return ""
    if len(text) <= limit:
        return text
    return text[:limit] + "\n[... source truncated for prompt ...]"


def _chunk_sample_source(text, share, chunk_size=2500):
    """Return a representative sample of ``text`` that fits inside ``share``
    characters by stitching together evenly spaced chunks.

    Mirrors the chunking pattern used by ``analyze_writing_quality_chunked``
    in ``utils.py``: split the source into fixed-size chunks, then keep the
    first chunk, the last chunk, and as many evenly spaced middle chunks as
    the budget allows. This preserves narrative anchors (opening/closing) and
    distributed evidence rather than dropping everything past a head cut-off.
    """
    if not text:
        return ""
    if len(text) <= share:
        return text

    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
    if len(chunks) <= 2:
        return _truncate(text, share)

    # Always keep first + last; pick evenly spaced middle chunks until the
    # share budget is consumed.
    selected_indices = {0, len(chunks) - 1}
    middle = list(range(1, len(chunks) - 1))
    if middle:
        budget_left = share - len(chunks[0]) - len(chunks[-1])
        # Sample evenly across the middle, prioritising spread over depth.
        step = max(1, len(middle) // max(1, budget_left // chunk_size))
        for j in range(0, len(middle), step):
            candidate = middle[j]
            tentative = sum(len(chunks[i]) for i in selected_indices | {candidate})
            if tentative > share:
                break
            selected_indices.add(candidate)

    ordered = sorted(selected_indices)
    pieces = []
    last_idx = -1
    for i in ordered:
        if i != last_idx + 1 and pieces:
            pieces.append(f"\n[... omitted intervening text from chunks "
                          f"{last_idx + 2}-{i} of {len(chunks)} ...]\n")
        pieces.append(chunks[i])
        last_idx = i
    sampled = "".join(pieces)
    if len(sampled) > share:
        sampled = sampled[:share] + "\n[... sampled extract truncated to fit prompt budget ...]"
    return sampled


def _format_sources_for_prompt(sources, char_budget=24000):
    """Concatenate the uploaded source items into a compact, labelled block.

    A simple proportional budget is applied so a single huge source does not
    crowd out the others. Long sources are sampled in chunks (head + spaced
    middles + tail) instead of being truncated at a hard head cut-off, so
    evidence from the entire document still reaches the prompt.
    """
    if not sources:
        return "(no source material uploaded yet)"

    total_chars = sum(len(s.get("text", "")) for s in sources) or 1
    blocks = []
    for idx, s in enumerate(sources, 1):
        text = s.get("text", "")
        share = max(800, int(char_budget * (len(text) / total_chars)))
        snippet = _chunk_sample_source(text, share)
        header = (
            f"--- SOURCE {idx} ---\n"
            f"Type: {s.get('type', 'unspecified')}\n"
            f"Speaker/Author: {s.get('speaker', 'unspecified')}\n"
            f"Date: {s.get('date', 'unspecified')}\n"
            f"Title/Filename: {s.get('name', 'unspecified')}\n"
            f"Original length: {len(text):,} chars; sample shown: {len(snippet):,} chars\n"
        )
        blocks.append(header + snippet)
    return "\n\n".join(blocks)


def _format_intake_for_prompt(metadata, intent):
    md = metadata or {}
    it = intent or {}
    pieces = [
        f"Case title: {md.get('title', '')}",
        f"Author(s): {md.get('authors', '')}",
        f"Initiative / programme: {md.get('initiative', '')}",
        f"Country / region: {md.get('region', '')}",
        f"Sector / theme: {md.get('sector', '')}",
        f"Primary protagonist: {md.get('protagonist', '')}",
        f"Key stakeholders: {md.get('stakeholders', '')}",
        f"Time period covered: {md.get('time_period', '')}",
        f"Intended audience: {md.get('audience', '')}",
        f"Target word count: {md.get('word_target', '')}",
        f"Competency alignment requested: {md.get('competencies', 'No')}",
        f"Teaching note requested: {'Yes' if md.get('teaching_note') else 'No'}",
        "",
        "Author's narrative intent:",
        f"  Central problem: {it.get('central_problem', '')}",
        f"  Dilemma / tension: {it.get('dilemma', '')}",
        f"  What went well: {it.get('went_well', '')}",
        f"  What did not go as expected: {it.get('did_not_go', '')}",
        f"  One key lesson for the reader: {it.get('key_lesson', '')}",
        f"  Preferred opening scene: {it.get('opening_scene', '')}",
    ]
    return "\n".join(pieces)


class GeneratorAPIError(RuntimeError):
    """Raised when an underlying AI call failed or returned an unusable
    response. Surfaced to the UI so the wizard never silently produces
    empty artefacts."""


def _call_json(prompt, temperature=0.1, seed=42):
    """Wrapper around call_openai_api that always asks for JSON.

    Raises ``GeneratorAPIError`` on transport failure or on the documented
    fallback shape returned by ``call_openai_api`` when JSON parsing fails.
    """
    full_prompt = MASTER_SYSTEM_PROMPT + "\n\n" + prompt
    result = call_openai_api(
        full_prompt,
        response_format="json_object",
        temperature=temperature,
        seed=seed,
    )
    # Outer-exception fallback in call_openai_api returns a string starting
    # with "Error calling OpenAI API:" — surface it as a hard error.
    if isinstance(result, str):
        raise GeneratorAPIError(result)
    if not isinstance(result, dict):
        raise GeneratorAPIError(
            f"Unexpected response type {type(result).__name__} from AI call"
        )
    # JSON-parse-failure fallback shape.
    if (
        result.get("reasoning") == "Failed to parse response from AI model"
        and result.get("document_reference") == "Error in analysis"
    ):
        raise GeneratorAPIError(
            "AI response could not be parsed as JSON. Please retry the step."
        )
    return result


def _call_text(prompt, temperature=0.1, seed=42):
    """Wrapper around call_openai_api that returns plain text.

    Raises ``GeneratorAPIError`` on transport failure.
    """
    full_prompt = MASTER_SYSTEM_PROMPT + "\n\n" + prompt
    result = call_openai_api(full_prompt, temperature=temperature, seed=seed)
    if not isinstance(result, str):
        raise GeneratorAPIError(
            f"Unexpected response type {type(result).__name__} from AI call"
        )
    if result.startswith("Error calling OpenAI API:"):
        raise GeneratorAPIError(result)
    return result


# ---------------------------------------------------------------------------
# Stage 2 — Source processing
# ---------------------------------------------------------------------------

def build_source_inventory(sources):
    """Prompt 2.1 — produce a per-source inventory of what each source contains."""
    sources_block = _format_sources_for_prompt(sources, char_budget=28000)

    prompt = f"""PROMPT 2.1 — Source Inventory

Read the source material below. Do NOT generate case content. Produce an
inventory entry for every source, listing only what the source actually
contains. Do not invent material.

=== SOURCE MATERIAL ===
{sources_block}

Return JSON with this exact shape:
{{
  "inventory": [
    {{
      "source_index": 1,
      "source_type": "interview transcript | report | news article | website | author notes | other",
      "speaker_or_author": "name and role",
      "date": "as known, otherwise 'unspecified'",
      "key_facts": ["short factual statement", "..."],
      "direct_quotes": [
        {{"quote": "verbatim quote", "speaker": "name and role"}}
      ],
      "events": [
        {{"approx_date": "YYYY or YYYY-MM", "event": "what happened", "actor": "who"}}
      ],
      "unsupported_claims": ["any claim flagged as not supported by evidence in the source"]
    }}
  ]
}}

Rules:
- direct_quotes must appear verbatim in the source.
- Leave any unknown field as the string "unspecified".
- If a source has no content of a given type, return an empty list.
"""
    return _call_json(prompt)


def build_chronology(sources, intake_metadata, intake_intent):
    """Prompt 2.2 — chronological timeline of events."""
    sources_block = _format_sources_for_prompt(sources, char_budget=28000)
    intake_block = _format_intake_for_prompt(intake_metadata, intake_intent)

    prompt = f"""PROMPT 2.2 — Chronology Extraction

Using only the source material, construct a chronological timeline of events
relevant to this case.

=== INTAKE ===
{intake_block}

=== SOURCE MATERIAL ===
{sources_block}

Return JSON:
{{
  "chronology": [
    {{
      "approx_date": "YYYY, YYYY-MM or 'unspecified'",
      "event": "what happened, in plain past tense",
      "actor": "name and role of the person responsible if stated, else 'unspecified'",
      "outcome": "what resulted, if stated, else 'unspecified'",
      "source_index": [1, 2]
    }}
  ],
  "gaps": ["short statement of a place where the source material is silent or unclear"]
}}

Rules:
- Sort chronologically (earliest first).
- Do not invent dates or actors. Use 'unspecified' if unsure.
- Mark source_index as the list of source numbers that support the entry.
"""
    return _call_json(prompt)


def build_stakeholder_map(sources, intake_metadata):
    """Prompt 2.3 — stakeholder mapping."""
    sources_block = _format_sources_for_prompt(sources, char_budget=28000)

    prompt = f"""PROMPT 2.3 — Stakeholder Mapping

From the source material, identify every stakeholder mentioned, including the
protagonist and key supporting actors.

Primary protagonist (per author): {(intake_metadata or {}).get('protagonist', '')}
Other stakeholders named by author: {(intake_metadata or {}).get('stakeholders', '')}

=== SOURCE MATERIAL ===
{sources_block}

Return JSON:
{{
  "stakeholders": [
    {{
      "name": "person, role, or organisation",
      "designation": "role or designation",
      "side": "demand | supply | convening/facilitation | other",
      "perspective": "their stated perspective if available, else 'unspecified'",
      "tensions": "any tension or hesitation expressed",
      "positive_outcomes": "any positive outcome they attributed",
      "source_index": [1, 2]
    }}
  ]
}}

Rules:
- Name only stakeholders the source material actually identifies.
- Where a name is withheld in the source, preserve the anonymity (e.g. "an
  unnamed broker").
"""
    return _call_json(prompt)


def build_dilemma_statement(sources, intake_metadata, intake_intent, case_type):
    """Prompt 2.4 — central dilemma statement."""
    sources_block = _format_sources_for_prompt(sources, char_budget=20000)
    intake_block = _format_intake_for_prompt(intake_metadata, intake_intent)
    case_label = CASE_TYPES.get(case_type, {}).get("label", case_type)

    paragraph_guidance = (
        "Write two short paragraphs (about 150-220 words total)."
        if case_type != "caselet"
        else "Write no more than two sentences."
    )

    prompt = f"""PROMPT 2.4 — Dilemma Identification

Based on the intake answers and the source material, draft a clear statement
of the central dilemma at the heart of this case. {paragraph_guidance}

The dilemma statement must:
- State clearly what the decision-maker or organisation was trying to achieve.
- Describe the obstacle or tension that made the goal difficult.
- Indicate why this was a genuine dilemma rather than a problem with an
  obvious answer.
- NOT resolve the dilemma in the statement itself.

This statement will anchor a {case_label}, so be especially precise about the
moment of choice or pattern of tension.

=== INTAKE ===
{intake_block}

=== SOURCE MATERIAL ===
{sources_block}

Return JSON:
{{
  "dilemma_statement": "the drafted paragraph(s) as a single string",
  "key_tensions": ["short bullet stating each tension"],
  "supporting_sources": [1, 2]
}}
"""
    return _call_json(prompt)


def run_source_processing(sources, intake_metadata, intake_intent, case_type, max_workers=4):
    """Run prompts 2.1-2.4 in parallel and return a single processed payload."""
    results = {}
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        futures = {
            "inventory": ex.submit(build_source_inventory, sources),
            "chronology": ex.submit(build_chronology, sources, intake_metadata, intake_intent),
            "stakeholders": ex.submit(build_stakeholder_map, sources, intake_metadata),
            "dilemma": ex.submit(
                build_dilemma_statement, sources, intake_metadata, intake_intent, case_type
            ),
        }
        for key, fut in futures.items():
            try:
                results[key] = fut.result()
            except Exception as e:
                results[key] = {"error": str(e)}
    return results


# ---------------------------------------------------------------------------
# Stage 3 — Drafting
# ---------------------------------------------------------------------------

# Per-section drafting prompts. Each entry is a tuple
# (display_word_guide_text, body_of_prompt) keyed by case type and section id.

LESSON_DRAWING_PROMPTS = {
    "hook": (
        "150-250 words",
        """Write the HOOK section.

Requirements:
- Open with a specific moment, scene, or situation drawn directly from the
  source material. It must be factual, not invented.
- Give enough context for an unfamiliar reader to understand what is at stake.
- Past tense, British English. 150-250 words.
- End with a sentence that signals the transition to the broader narrative.

Must NOT:
- Begin with a rhetorical question or a bare statistic.
- Use the protagonist's name in the first sentence.
- Describe the protagonist in evaluative terms.
- Reveal the outcome of the initiative.

Preferred opening cue from the author: "{opening_scene}"
""",
    ),
    "introduction": (
        "10-15% of total word count",
        """Write the INTRODUCTION AND CONTEXT SETTING section.

Narrative continuity with the Hook (CRITICAL — applies only when the
hook has already been drafted):
- If a "HOOK (already drafted, do not repeat)" block appears further
  down, read it before you write and follow all of the rules below.
  If no such block is present, skip this continuity guidance and
  write a standalone introduction.
- Your FIRST sentence must pick up directly from the closing image or
  moment of the hook. Do NOT reset to a fresh topic, do NOT open with a
  generic statement about the sector, and do NOT open with the
  protagonist's name in isolation.
- Reuse at least one concrete noun, actor, place or setting from the
  hook in your first paragraph so the reader feels they are still in
  the same story.
- Do NOT repeat or paraphrase facts the hook has already stated. After
  the bridging opener, MOVE the reader outward from the hook's specific
  moment into the systemic / sectoral context.

Requirements:
- Establish the systemic or sectoral context.
- Introduce the primary protagonist by name, role and institutional affiliation.
  State their mandate at the time. Do not evaluate their character.
- Introduce the convening or sponsoring institution and explain why this
  initiative fell within its mandate.
- Describe the state of the sector before the initiative began: what was
  working, what was not, and what observations drew attention to the gap.
- Be factual and neutral. Do not call anything a 'crisis' or 'failure' unless
  the source material explicitly does so with attribution.
- Flag specialist or country-specific terms for footnotes using
  [FOOTNOTE NEEDED: term].
- British English, past tense.
""",
    ),
    "dilemma": (
        "10-12% of total word count",
        """Write the DEVELOPMENT CHALLENGE AND CENTRAL DILEMMA section. Use two parts.

Narrative continuity with the previous section (CRITICAL — applies only
when the previous section has already been drafted):
- If a "PREVIOUS SECTION ... (already drafted, do not repeat)" block
  appears further down, read it before you write and follow all of the
  rules below. If no such block is present, skip this continuity
  guidance and write a standalone section.
- Your FIRST sentence must bridge directly from the closing image,
  moment, actor or idea of the previous section. Do NOT reset to a
  fresh topic, do NOT open with a generic statement about the sector
  or the problem, and do NOT open with an abstract definition of the
  dilemma.
- Reuse at least one concrete noun, actor, place or moment from the
  previous section in your opening paragraph so the reader feels they
  are still inside the same story.
- Do NOT repeat or paraphrase facts the previous section has already
  stated. After the bridging opener, move the reader from that moment
  into the structural challenge below.

Part A — Development Challenge
Describe the structural or systemic problem the initiative was designed to
address. This is a pattern or condition, not a single event. Use sub-headings
if the challenge has more than one dimension.

Part B — Central Dilemma
Expand the confirmed dilemma statement (provided below) into two to three
paragraphs. State it clearly enough that a reader understands why it was
genuinely difficult. Do not resolve it. Close with a sentence signalling that
the case will trace how the protagonist navigated this tension.

Confirmed dilemma statement:
{dilemma_statement}

Requirements:
- Do not use 'the solution was' or any equivalent phrase that prematurely
  resolves the dilemma.
- Include at least one direct quote from the source material attributed to a
  named speaker.
- British English, past tense.
""",
    ),
    "delivery_challenges": (
        "10-12% of total word count",
        """Write the DELIVERY CHALLENGES section.

Narrative continuity with the previous section (CRITICAL — applies only
when the previous section has already been drafted):
- If a "PREVIOUS SECTION ... (already drafted, do not repeat)" block
  appears further down, read it before you write and follow all of the
  rules below. If no such block is present, skip this continuity
  guidance and write a standalone section.
- Your FIRST sentence must bridge directly from the closing image,
  moment, actor or tension of the previous section (typically the
  Central Dilemma). Do NOT open with a bare sub-heading, do NOT open
  with a generic statement such as "Several challenges arose", and do
  NOT open with a list.
- Reuse at least one concrete noun, actor, place or moment from the
  previous section in the opening sentence so the reader stays inside
  the same story before the first sub-heading.
- Do NOT repeat or paraphrase facts the previous section has already
  stated. After the bridging opener, move the reader into the
  challenges below.

For each delivery challenge present in the source material:
- Name it with a sub-heading.
- Describe what it was and why it arose.
- Describe how it was addressed, if the source material says so. If it
  remained unresolved, say so.
- Attribute descriptions to a named source or identified event wherever
  possible.

Requirements:
- Present challenges as facts, not as criticism of individuals or institutions.
- Cover funding, trust, capacity, institutional resistance and sustainability
  challenges where they appear in the source material.
- Use concrete details: amounts, timelines, institution names where available.
- British English, past tense.
""",
    ),
    "implementation": (
        "35-40% of total word count — the longest section",
        """Write the CHRONOLOGICAL IMPLEMENTATION section.

Narrative continuity with the previous section (CRITICAL — applies only
when the previous section has already been drafted):
- If a "PREVIOUS SECTION ... (already drafted, do not repeat)" block
  appears further down, read it before you write and follow all of the
  rules below. If no such block is present, skip this continuity
  guidance and write a standalone section.
- Your FIRST sentence must bridge directly from the closing image,
  moment, actor or unresolved tension of the previous section
  (typically the Delivery Challenges). Do NOT open with a bare phase
  sub-heading, do NOT open with a generic statement such as "The
  initiative unfolded in several phases", and do NOT begin with a
  date in isolation.
- Reuse at least one concrete noun, actor, place or moment from the
  previous section in the opening sentence so the reader stays inside
  the same story before the first phase sub-heading.
- Do NOT repeat or paraphrase facts the previous section has already
  stated. After the bridging opener, move into the first chronological
  phase.

This is the longest section of the case. Use the confirmed chronology as the
backbone and group events into chronological sub-sections (for example,
'The early phase: 2020-2021', 'Scaling up: 2022', 'Consolidation: 2023').

For each phase:
- Describe what happened, who made which decision and what the outcome was.
- Incorporate stakeholder perspectives at the relevant point. Attribute every
  perspective to a named source.
- Describe two or three of the most instructive collaborations, projects or
  outcomes in enough detail for the reader to understand why they matter.
- If something did not go as expected, say so plainly.

Requirements:
- Strict chronological order. Signal any shift in time explicitly.
- Use sub-headings for each major phase.
- Maintain neutral tone throughout.
- Include at least three direct quotes from named sources.
- Flag specialist terms for footnotes using [FOOTNOTE NEEDED: term].
- British English, past tense.

Confirmed chronology to use as the backbone:
{chronology_text}
""",
    ),
    "lessons": (
        "12-15% of total word count",
        """Write the LESSONS LEARNT section.

Narrative continuity with the previous section (CRITICAL — applies only
when the previous section has already been drafted):
- If a "PREVIOUS SECTION ... (already drafted, do not repeat)" block
  appears further down, read it before you write and follow all of the
  rules below. If no such block is present, skip this continuity
  guidance and write a standalone section.
- Your FIRST sentence (before any lesson sub-heading) must bridge
  directly from the closing image, moment, actor or phase of the
  previous section (typically the final phase of Implementation). Do
  NOT open with a bare lesson sub-heading, do NOT open with a generic
  framing such as "Several lessons emerge from this case", and do NOT
  jump straight into prescriptions.
- Reuse at least one concrete noun, actor, place or moment from the
  previous section in that opening sentence so the reader feels the
  lessons grow out of the story they have just read.
- Do NOT repeat or paraphrase facts the previous section has already
  stated. After the bridging opener, move into the first lesson.

For each lesson:
- Give it a clear sub-heading stating the lesson in plain language.
- Write two to four paragraphs explaining: what happened in the case that
  demonstrates it, why it matters beyond this case, and what a similar
  organisation might do differently.
- Include a direct quote from a named source where it captures the lesson
  concisely.
- Where a lesson is relevant to a specified competency, note this within the
  paragraph (without naming the competency framework).

Requirements:
- Present between four and seven lessons.
- Do NOT introduce any new facts, events or sources that have not appeared
  earlier in the case.
- Present lessons as observations drawn from evidence, not as prescriptions.
- Acknowledge where a lesson points to something that remained unresolved.
- British English. General principles within the lesson may be in the present
  tense; descriptive material remains in past tense.
""",
    ),
    "closing": (
        "150-250 words",
        """Write the CLOSING section.

Narrative continuity with the previous section (CRITICAL — applies only
when the previous section has already been drafted):
- If a "PREVIOUS SECTION ... (already drafted, do not repeat)" block
  appears further down, read it before you write and follow all of the
  rules below. If no such block is present, skip this continuity
  guidance and write a standalone section.
- Your FIRST sentence must bridge directly from the closing image,
  moment, actor or lesson of the previous section (typically the
  final Lesson Learnt). Do NOT reset to a generic summary framing such
  as "In conclusion" or "Looking back", and do NOT restate the
  initiative's name in isolation.
- Reuse at least one concrete noun, actor, place or moment from the
  previous section in your opening sentence so the closing feels like
  the same story coming to rest.
- Do NOT repeat or paraphrase facts the previous section has already
  stated. After the bridging opener, return to the central dilemma as
  required below.

The Closing must:
- Return to the central dilemma in the first paragraph and describe, in
  qualified terms, how the initiative navigated it. Do not present the outcome
  as a complete resolution.
- In a second paragraph, state what the initiative represented at the level of
  the system or sector.
- Close with a sentence signalling the most important unresolved question or
  the condition that would determine whether the gains made would last.

Must NOT:
- Use phrases such as 'a shining example', 'a model for the world', or
  'transforming the landscape'.
- Introduce new information.
- End with a quotation.

British English, past tense. 150-250 words.
""",
    ),
}

DECISION_FORCING_PROMPTS = {
    "executive_summary": (
        "~200 words",
        """Write the EXECUTIVE SUMMARY of this decision-forcing case.

Requirements:
- About 200 words.
- Introduce the protagonist by name, role and institution.
- Describe the situation they faced at the point of decision.
- State the nature of the decision they had to make WITHOUT revealing what
  they actually decided.
- Past tense, British English.
- Engaging and clear; should make the reader want to read on.

Must NOT:
- Include any detail about the outcome or the choice the protagonist made.
- Use evaluative language about the protagonist's character or capabilities.
""",
    ),
    "introduction": (
        "~1 page (350-500 words)",
        """Write the INTRODUCTION of this decision-forcing case.

Narrative continuity with the Executive Summary (CRITICAL — applies
only when the executive summary has already been drafted):
- If an "EXECUTIVE SUMMARY (already drafted, do not repeat)" block
  appears further down, read it before you write and follow all of
  the rules below. If no such block is present, skip this continuity
  guidance and write a standalone introduction.
- Your FIRST sentence must pick up directly from the closing image or
  moment of the executive summary. Do NOT reset to a fresh topic and
  do NOT open with a generic scene-setter unrelated to it.
- Reuse at least one concrete noun, actor, place or setting from the
  executive summary in your first paragraph so the reader stays
  inside the same story.
- Do NOT repeat or paraphrase facts the executive summary has already
  stated. After the bridging opener, widen the lens to the
  institutional setting, protagonist's role, and the nature of the
  decision.

Requirements:
- Set the scene: where, when and in what institutional context the case takes
  place.
- Introduce the protagonist by name, role and the scope of their authority.
- State briefly what kind of decision the case concerns, WITHOUT revealing the
  options or outcome.
- About one page (350-500 words).
- British English, past tense.
""",
    ),
    "background": (
        "2-3 pages (700-1100 words)",
        """Write the BACKGROUND section.

Narrative continuity with the previous section (CRITICAL — applies only
when the previous section has already been drafted):
- If a "PREVIOUS SECTION ... (already drafted, do not repeat)" block
  appears further down, read it before you write and follow all of the
  rules below. If no such block is present, skip this continuity
  guidance and write a standalone section.
- Your FIRST sentence must bridge directly from the closing image,
  moment or actor of the previous section (typically the
  Introduction). Do NOT reset to a textbook-style opening on the
  sector, do NOT begin with a date in isolation, and do NOT restate
  the protagonist's name and title.
- Reuse at least one concrete noun, actor, place or moment from the
  previous section in your opening paragraph so the reader stays
  inside the same story.
- Do NOT repeat or paraphrase facts the previous section has already
  stated. After the bridging opener, widen the lens into the sector
  and historical context below.

Requirements:
- Describe the sector, system or institutional context. Draw on the confirmed
  source material.
- Trace the history of the situation: how did it develop? What conditions made
  a decision necessary?
- Identify the key stakeholders and describe their interests, constraints and
  relationships to the protagonist.
- Present only the information the protagonist would have had at the time of
  the decision. No hindsight bias.
- 2-3 pages (700-1100 words).
- British English, past tense.
- Flag specialist terms for footnotes using [FOOTNOTE NEEDED: term].
""",
    ),
    "narrative": (
        "3-5 pages (1100-1800 words)",
        """Write the NARRATIVE section.

Narrative continuity with the previous section (CRITICAL — applies only
when the previous section has already been drafted):
- If a "PREVIOUS SECTION ... (already drafted, do not repeat)" block
  appears further down, read it before you write and follow all of the
  rules below. If no such block is present, skip this continuity
  guidance and write a standalone section.
- Your FIRST sentence must bridge directly from the closing image,
  moment or actor of the previous section (typically the Background).
  Do NOT begin with a date in isolation, do NOT restate the
  protagonist's name and title, and do NOT open with a generic
  framing such as "Events unfolded as follows".
- Reuse at least one concrete noun, actor, place or moment from the
  previous section in your opening paragraph so the reader stays
  inside the same story before the first event in the chronology.
- Do NOT repeat or paraphrase facts the previous section has already
  stated. After the bridging opener, move into the first event in the
  chronology.

Requirements:
- Trace the sequence of events that led directly to the decision point. Use
  the confirmed chronology below as the backbone.
- Show how the pressures, constraints and stakeholder dynamics accumulated to
  make a decision unavoidable.
- Include the protagonist's perspective from the source material, attributed
  to named sources.
- Build the reader's understanding of what made the decision difficult,
  WITHOUT steering them towards a particular answer.
- Include at least two direct quotes from named sources.
- 3-5 pages (1100-1800 words).
- British English, past tense.

Must NOT:
- Reveal the protagonist's actual choice.
- Frame any option as obviously correct or wrong.
- Include outcomes that occurred after the decision point.

Confirmed chronology:
{chronology_text}
""",
    ),
    "dilemma": (
        "~300 words",
        """Write the DILEMMA OR PUZZLE section. About 300 words.

Narrative continuity with the previous section (CRITICAL — applies only
when the previous section has already been drafted):
- If a "PREVIOUS SECTION ... (already drafted, do not repeat)" block
  appears further down, read it before you write and follow all of the
  rules below. If no such block is present, skip this continuity
  guidance and write a standalone section.
- Your FIRST sentence must bridge directly from the closing image,
  moment or actor of the previous section (typically the Narrative).
  Do NOT reset to an abstract framing of the decision, do NOT begin
  with a generic line such as "The protagonist now faced a difficult
  choice", and do NOT restate the protagonist's name and title.
- Reuse at least one concrete noun, actor, place or moment from the
  previous section in your opening sentence so the reader feels the
  decision arising from the events they have just read.
- Do NOT repeat or paraphrase facts the previous section has already
  stated. After the bridging opener, state the decision precisely as
  required below.

Requirements:
- State the decision the protagonist faced, clearly and precisely, at the
  exact moment in time where the case stops.
- Describe at least two distinct options available.
- For each option, briefly note its main advantages and the main risks or
  trade-offs it carries.
- Describe the constraints the protagonist faced: time, resources,
  institutional authority, political environment, stakeholder relationships.
- End with a direct statement addressed to the reader: "In this situation,
  what would you do?"

Must NOT:
- Include the protagonist's actual decision.
- Use language that implies one option is the right answer.
- Introduce information not already established in Background or Narrative.

Confirmed dilemma statement:
{dilemma_statement}
""",
    ),
    "epilogue": (
        "Optional — up to 300 words",
        """Write a brief EPILOGUE for this decision-forcing case.

Narrative continuity with the previous section (CRITICAL — applies only
when the previous section has already been drafted):
- If a "PREVIOUS SECTION ... (already drafted, do not repeat)" block
  appears further down, read it before you write and follow all of the
  rules below. If no such block is present, skip this continuity
  guidance and write a standalone section.
- Your FIRST sentence must bridge directly from the decision question
  posed at the end of the previous section (the Dilemma or Puzzle)
  and reveal what the protagonist actually decided. Do NOT reset to a
  generic framing such as "In the event, ...", and do NOT restate the
  protagonist's name and title in isolation.
- Reuse at least one concrete noun, actor, place or option from the
  previous section in your opening sentence so the epilogue answers
  the question the reader was just asked.
- Do NOT repeat or paraphrase facts the previous section has already
  stated. After the bridging opener, describe what followed.

Requirements:
- State what the protagonist actually decided.
- Describe what happened as a result, based on the source material.
- Factual and neutral; do not evaluate whether the decision was correct.
- No more than 300 words.
- Note any outcomes that were not anticipated at the time of the decision.
- British English, past tense.
""",
    ),
    "study_questions": (
        "3-5 questions",
        """Write 3-5 STUDY QUESTIONS for this decision-forcing case.

Requirements:
- Require the reader to ANALYSE the situation, not merely recall facts.
- Include at least one question asking the reader to choose and defend a
  course of action.
- Include at least one question asking the reader to take the perspective of
  a specific stakeholder other than the protagonist.
- Include at least one 'what if' question: what would have changed if a
  specific condition had been different?
- Phrase as open questions that cannot be answered in a single word.

Must NOT:
- Have a single correct answer that can be read directly from the case.
- Repeat the same analytical angle as one another.

Return as a numbered list (1., 2., 3., ...).
""",
    ),
}


def _build_drafting_context(intake_metadata, intake_intent, sources, processed):
    """Common shared context block for every drafting prompt."""
    sources_block = _format_sources_for_prompt(sources, char_budget=20000)
    intake_block = _format_intake_for_prompt(intake_metadata, intake_intent)

    chronology = processed.get("chronology", {}).get("chronology", []) if processed else []
    chronology_lines = [
        f"- {c.get('approx_date', 'unspecified')}: {c.get('event', '')} "
        f"(actor: {c.get('actor', 'unspecified')}; outcome: {c.get('outcome', 'unspecified')})"
        for c in chronology
    ]
    chronology_text = "\n".join(chronology_lines) if chronology_lines else "(no chronology yet)"

    stakeholders = processed.get("stakeholders", {}).get("stakeholders", []) if processed else []
    stakeholder_lines = [
        f"- {s.get('name', '')} ({s.get('designation', '')}): {s.get('perspective', '')}"
        for s in stakeholders
    ]
    stakeholder_text = "\n".join(stakeholder_lines) if stakeholder_lines else "(no stakeholders yet)"

    dilemma_statement = (
        processed.get("dilemma", {}).get("dilemma_statement", "")
        if processed
        else ""
    )

    return {
        "sources_block": sources_block,
        "intake_block": intake_block,
        "chronology_text": chronology_text,
        "stakeholder_text": stakeholder_text,
        "dilemma_statement": dilemma_statement or (intake_intent or {}).get("dilemma", ""),
    }


def draft_case_section(
    case_type,
    section_id,
    intake_metadata,
    intake_intent,
    sources,
    processed,
    previous_sections=None,
    author_feedback="",
):
    """Draft a single section of a Lesson-Drawing or Decision-Forcing case."""
    previous_sections = previous_sections or {}
    ctx = _build_drafting_context(intake_metadata, intake_intent, sources, processed)

    if case_type == "lesson_drawing":
        section_prompts = LESSON_DRAWING_PROMPTS
    elif case_type == "decision_forcing":
        section_prompts = DECISION_FORCING_PROMPTS
    else:
        raise ValueError(f"draft_case_section does not support case_type={case_type!r}")

    if section_id not in section_prompts:
        raise ValueError(f"Unknown section {section_id!r} for {case_type}")

    word_guide, body = section_prompts[section_id]
    body = body.format(
        chronology_text=ctx["chronology_text"],
        stakeholder_text=ctx["stakeholder_text"],
        dilemma_statement=ctx["dilemma_statement"],
        opening_scene=(intake_intent or {}).get("opening_scene", ""),
    )

    prior_block = ""
    hook_block = ""
    # Surface the IMMEDIATELY PREVIOUS already-drafted section in its own
    # clearly-labelled "bridge" block so the model uses it as the narrative
    # bridge for this section. This applies to every non-opening section
    # (introduction onwards), not just the introduction. To avoid prompt
    # bloat and duplicate content, that previous section is excluded from
    # the generic prior_block below.
    #
    # For the introduction specifically, the immediately-previous section
    # is the opener (hook / executive_summary) and we keep the original
    # HOOK / EXECUTIVE SUMMARY labelling so the existing intro-prompt
    # continuity rules still match.
    section_order = [sid for sid, _ in CASE_TYPES[case_type]["sections"]]
    bridge_sid = None
    if section_id in section_order:
        idx = section_order.index(section_id)
        for prev_idx in range(idx - 1, -1, -1):
            candidate = section_order[prev_idx]
            if (previous_sections.get(candidate) or {}).get("text"):
                bridge_sid = candidate
                break

    if previous_sections:
        joined = []
        for sid, sname in CASE_TYPES[case_type]["sections"]:
            if sid == bridge_sid:
                continue
            if sid in previous_sections and previous_sections[sid].get("text"):
                joined.append(f"### {sname}\n{previous_sections[sid]['text']}")
        if joined:
            prior_block = "\n\n=== ALREADY-DRAFTED EARLIER SECTIONS ===\n" + "\n\n".join(joined)

        if bridge_sid:
            bridge_text = ((previous_sections.get(bridge_sid) or {}).get("text") or "").strip()
            if bridge_text:
                bridge_name = dict(CASE_TYPES[case_type]["sections"]).get(bridge_sid, bridge_sid)
                # Preserve the original HOOK / EXECUTIVE SUMMARY labels for
                # the introduction so the intro prompt's existing continuity
                # rules (which reference those exact labels) still match.
                if section_id == "introduction" and bridge_sid in ("hook", "executive_summary"):
                    bridge_label = "HOOK" if bridge_sid == "hook" else "EXECUTIVE SUMMARY"
                    hook_block = (
                        f"\n\n=== {bridge_label} (already drafted, do not repeat) ===\n"
                        f"{bridge_text}\n"
                        f"=== END {bridge_label} ===\n"
                        "Your introduction's first sentence MUST bridge from the "
                        "closing image / moment of the block above, reuse at least "
                        "one concrete noun / actor / place from it, and then widen "
                        "out. Do not restate its facts."
                    )
                else:
                    bridge_label = f"PREVIOUS SECTION: {bridge_name.upper()}"
                    hook_block = (
                        f"\n\n=== {bridge_label} (already drafted, do not repeat) ===\n"
                        f"{bridge_text}\n"
                        f"=== END {bridge_label} ===\n"
                        "Your section's FIRST sentence MUST bridge from the "
                        "closing image / moment / actor of the block above, "
                        "reuse at least one concrete noun / actor / place / "
                        "moment from it, and then move into the content this "
                        "section requires. Do not restate its facts and do "
                        "not open with a bare sub-heading or generic framing."
                    )

    feedback_block = ""
    if author_feedback:
        feedback_block = (
            "\n\n=== AUTHOR FEEDBACK FOR THIS REGENERATION ===\n"
            f"{author_feedback.strip()}\n"
            "Please incorporate this feedback while still respecting the rules above."
        )

    full_prompt = f"""You are drafting a {CASE_TYPES[case_type]['label']} for the
Amrit Gyaan Kosh repository. Word guide for this section: {word_guide}.

{body}

=== INTAKE ===
{ctx['intake_block']}

=== STAKEHOLDERS (confirmed) ===
{ctx['stakeholder_text']}

=== SOURCE MATERIAL ===
{ctx['sources_block']}
{prior_block}{hook_block}{feedback_block}

Now write the section. Return ONLY the section body text — no headings beyond
sub-headings inside the section, no commentary, no JSON, no markdown fences.
"""
    return _call_text(full_prompt)


def draft_caselet(intake_metadata, intake_intent, sources, processed, author_feedback=""):
    """Draft an entire Caselet in a single call (Prompt 3C.1)."""
    ctx = _build_drafting_context(intake_metadata, intake_intent, sources, processed)
    feedback_block = ""
    if author_feedback:
        feedback_block = (
            "\n\n=== AUTHOR FEEDBACK FOR THIS REGENERATION ===\n"
            f"{author_feedback.strip()}\n"
        )

    prompt = f"""PROMPT 3C.1 — Full Caselet Draft

Write a CASELET based on the source material and intake information below.
The caselet must NOT exceed 2,000 words, excluding footnotes and references.

Follow this exact structure with these sub-headings:

## Context (150-250 words)
Introduce the setting, the institution and the protagonist. State what they
were responsible for and what the general situation was. Do not provide a
detailed history.

## The Situation (400-700 words)
Describe the specific challenge, decision point or practice that is the focus
of the caselet. Be concrete and specific. Draw only on the source material.
Include one direct quote from a named source if available. Flag any specialist
terms for footnotes using [FOOTNOTE NEEDED: term].

## The Response or Outcome (250-400 words)
Describe what was decided or done, and what resulted. If this caselet is being
left open as a decision-forcing exercise, stop before the outcome and state
clearly: "At this point, [name] had to decide [state the decision]."

## Discussion Questions
Write one or two open questions answerable through reflection on the caselet
without additional research. At least one question should ask what the reader
would have done.

All writing in British English and past tense. No unnecessary adjectives. No
evaluative language about individuals.

=== INTAKE ===
{ctx['intake_block']}

=== STAKEHOLDERS (confirmed) ===
{ctx['stakeholder_text']}

=== SOURCE MATERIAL ===
{ctx['sources_block']}

Confirmed dilemma statement:
{ctx['dilemma_statement']}
{feedback_block}

Return ONLY the caselet body, with the sub-headings shown above. No JSON, no
markdown fences, no commentary.
"""
    return _call_text(prompt)


# ---------------------------------------------------------------------------
# Stage 4 — Compliance passes
# ---------------------------------------------------------------------------

def _assemble_full_draft(case_type, draft_sections):
    """Concatenate per-section text into a single document for compliance review."""
    if case_type == "caselet":
        return (draft_sections or {}).get("caselet", {}).get("text", "")

    parts = []
    for sid, sname in CASE_TYPES[case_type]["sections"]:
        sec = (draft_sections or {}).get(sid)
        if sec and sec.get("text"):
            parts.append(f"## {sname}\n\n{sec['text']}")
    return "\n\n".join(parts)


def run_compliance_passes(case_type, draft_sections, intake_metadata, sources):
    """Run the cross-cutting style/structure/sensitivity/word-count passes
    described in Prompt Groups 4-9 and return a unified findings list."""
    full_text = _assemble_full_draft(case_type, draft_sections)
    if not full_text.strip():
        return {"findings": [], "word_count": 0, "target": 0, "abbreviations": [], "footnotes": [], "references": []}

    target = int((intake_metadata or {}).get("word_target") or
                 CASE_TYPES.get(case_type, {}).get("default_target", 0))
    case_label = CASE_TYPES[case_type]["label"]

    sources_block = _format_sources_for_prompt(sources, char_budget=16000)

    style_prompt = f"""You are running automated style and structure checks
against a {case_label}. The full draft is below.

Run ALL of the following checks and return a single consolidated JSON
findings list:

PROMPT GROUP 4 — Language and style
- 4.1 British English: flag US spellings/idioms with the BR equivalent.
- 4.2 Tense: flag sentences in present/future tense outside allowed exceptions
  (direct quotes, general principles in a Lessons section).
- 4.3 Adjective audit: flag evaluative adjectives/adverbs not supported by
  data, especially: significant, major, important, critical, key, innovative,
  unique, pioneering, transformative, robust, comprehensive, holistic,
  seamless, impactful, stellar, extraordinary, remarkable, relentless,
  dedicated, passionate.
- 4.4 Em-dashes & sentence length: flag every em-dash, every sentence over
  45 words, every paragraph that is a single sentence (except in Hook,
  Caselet Context or Closing), and three or more consecutive paragraphs
  starting with the same word.
- 4.5 Neutrality: flag hero narrative, promotional language, political
  sensitivity, implicit criticism without attributed source, and blame
  attribution without an attributed source.

PROMPT GROUP 7 — Structure
- 7.1 Section sequence and presence (against the canonical {case_label}
  ordering).
- 7.2 Coherence/transitions between major sections.

PROMPT GROUP 8 — Political, cultural and confidentiality
- Political party / electoral / cross-country governance judgements.
- Gendered language defaulting to 'he', or unnecessary gender, regional,
  caste, religious or community references.
- Commercially sensitive or personal information about named individuals.

PROMPT GROUP 9 — Word count & formatting standards
- 9.2 formatting: heading consistency, double quotation marks, footnote
  marker placement, abbreviations rule on first use, paragraph spacing, no
  one-sentence paragraphs outside the allowed exceptions.

Return JSON:
{{
  "findings": [
    {{
      "type": "British English | Tense | Adjective | Em-dash | Sentence length | Single-sentence paragraph | Repetition | Hero narrative | Promotional | Political | Implicit criticism | Blame | Section sequence | Coherence | Gendered language | Identity reference | Confidentiality | Formatting | Other",
      "section": "section heading where the issue was found, or 'Whole document'",
      "original_text": "short verbatim excerpt (max 25 words)",
      "suggestion": "concrete suggested revision",
      "severity": "High | Medium | Low",
      "explanation": "one short sentence"
    }}
  ]
}}

Be conservative. Only flag genuine issues. If the draft is clean for a check,
do not invent findings for it.

=== DRAFT ===
{full_text[:18000]}
"""

    sensitivity_prompt = f"""PROMPT 6 / 8 — Claim, quote and sensitivity audit
for a {case_label}.

Compare the DRAFT against the SOURCE MATERIAL and report:
- Specific factual claims (numbers, dates, named people/institutions/programmes,
  events, attributed quotes) for which no source can be identified — flag as
  'Untraceable claim'.
- Direct quotations that do not appear verbatim in the sources — flag as
  'Quote mismatch'.
- Numerical figures whose value or unit differs from the source — flag as
  'Numerical mismatch'.
- Any pre-publication sensitivity (pending decisions, ongoing legal matters,
  unannounced policies).

Return JSON with the SAME shape as the style pass:
{{
  "findings": [
    {{
      "type": "Untraceable claim | Quote mismatch | Numerical mismatch | Pre-publication sensitivity",
      "section": "section heading or 'Whole document'",
      "original_text": "short verbatim excerpt (max 25 words)",
      "suggestion": "what the author should do (cite source / verify / remove / generalise)",
      "severity": "High | Medium | Low",
      "explanation": "one sentence"
    }}
  ]
}}

=== SOURCE MATERIAL ===
{sources_block}

=== DRAFT ===
{full_text[:14000]}
"""

    abbrev_prompt = f"""PROMPT 5 — Abbreviations, footnotes and references for
a {case_label}.

Run three checks:

1. Abbreviation audit: list every multi-word term in the draft that has, or
   could have, an abbreviation. Count occurrences. If 3 or more, the
   abbreviation should be introduced in parentheses on first use; otherwise
   write the full term every time.

2. Footnote suggestions: identify country-specific or sector-specific terms a
   non-specialist international reader may not understand and that are not
   already explained in body. For each, draft a footnote of <=3 sentences.

3. Reference list: based on the SOURCE MATERIAL provided, build a Harvard-
   style reference list of every source actually used in the draft.

Return JSON:
{{
  "abbreviations": [
    {{"term": "full term", "abbreviation": "ABC", "occurrences": 4, "action": "introduce on first use" | "write in full each time"}}
  ],
  "footnotes": [
    {{"term": "term as it appears in the draft", "footnote_text": "<=3 sentence explanation"}}
  ],
  "references": [
    "Surname, Initial. (Year). Title. Source. Retrieved from URL."
  ]
}}

=== SOURCE MATERIAL ===
{sources_block}

=== DRAFT ===
{full_text[:14000]}
"""

    word_count = len(re.findall(r"\b\w+\b", full_text))

    findings = []
    abbreviations = []
    footnotes = []
    references = []

    with ThreadPoolExecutor(max_workers=3) as ex:
        f_style = ex.submit(_call_json, style_prompt)
        f_sense = ex.submit(_call_json, sensitivity_prompt)
        f_abbrev = ex.submit(_call_json, abbrev_prompt)

        try:
            style_res = f_style.result()
            for f in style_res.get("findings", []) or []:
                if isinstance(f, dict):
                    findings.append(f)
        except Exception as e:
            findings.append({"type": "Other", "section": "Whole document",
                             "original_text": "", "suggestion": "",
                             "severity": "Low",
                             "explanation": f"Style pass error: {e}"})

        try:
            sense_res = f_sense.result()
            for f in sense_res.get("findings", []) or []:
                if isinstance(f, dict):
                    findings.append(f)
        except Exception as e:
            findings.append({"type": "Other", "section": "Whole document",
                             "original_text": "", "suggestion": "",
                             "severity": "Low",
                             "explanation": f"Sensitivity pass error: {e}"})

        try:
            abbrev_res = f_abbrev.result()
            abbreviations = abbrev_res.get("abbreviations", []) or []
            footnotes = abbrev_res.get("footnotes", []) or []
            references = abbrev_res.get("references", []) or []
        except Exception:
            pass

    # Word count finding (Prompt 9.1)
    if target:
        delta = word_count - target
        pct = abs(delta) / target * 100 if target else 0
        if case_type == "caselet" and word_count > 2000:
            findings.append({
                "type": "Word count",
                "section": "Whole document",
                "original_text": f"{word_count} words",
                "suggestion": "Caselet ceiling is 2,000 words. Identify cuts.",
                "severity": "High",
                "explanation": f"Caselet exceeds the 2,000-word ceiling by {word_count - 2000} words.",
            })
        elif pct > 10:
            direction = "over" if delta > 0 else "under"
            findings.append({
                "type": "Word count",
                "section": "Whole document",
                "original_text": f"{word_count} words against target {target}",
                "suggestion": (
                    "Identify passages that repeat the same point or contain "
                    "detail beyond what is needed to establish the facts."
                    if direction == "over"
                    else "Incorporate additional evidence from the source material or stakeholder perspectives not yet included."
                ),
                "severity": "Medium",
                "explanation": f"Draft is {pct:.1f}% {direction} the target word count.",
            })

    return {
        "findings": findings,
        "word_count": word_count,
        "target": target,
        "abbreviations": abbreviations,
        "footnotes": footnotes,
        "references": references,
    }


# ---------------------------------------------------------------------------
# Stage 5 — Optional teaching note (Prompt Group 12)
# ---------------------------------------------------------------------------

def generate_teaching_note(case_type, draft_sections, intake_metadata, sources, processed):
    """Generate teaching note sections 12.1-12.7 in a single structured call."""
    full_text = _assemble_full_draft(case_type, draft_sections)
    if not full_text.strip():
        return {}

    case_label = CASE_TYPES[case_type]["label"]
    audience = (intake_metadata or {}).get("audience", "")
    competencies = (intake_metadata or {}).get("competencies", "")
    sources_block = _format_sources_for_prompt(sources, char_budget=12000)

    prompt = f"""PROMPT GROUP 12 — Generate a TEACHING NOTE for a {case_label}.

You will produce all seven sections of the Amrit Gyaan Kosh teaching note
template in one structured response. Base every section on the confirmed case
draft below. Do not invent new facts.

Sections required (12.1-12.7):
1. Case Summary (<=400 words)
2. Learning Objectives (3-4, action-verb-led, derived from the case)
3. Teaching Strategy (recommended session format, suggested duration based on
   case type, discussion flow, facilitation notes, common participant errors)
4. Discussion Questions with Answers (each question + a detailed answer or
   answer guide; for open questions provide key analytical points, two or
   three alternative positions and a note on superficial responses; if a
   theoretical framework applies, describe it briefly and apply it to the
   case)
5. Assessment Methodology (one or two specific approaches linked to the
   Learning Objectives; specify length and assessment criteria)
6. Target Audience and Subject Domains (state target audience from intake;
   list subject domains; note any prior knowledge required)
7. Additional Teaching Resources (2-4 supplementary resources with title,
   source, one-sentence relevance and access note)

If competency alignment was requested, at least one Learning Objective should
correspond to that competency:
{competencies}

Intended audience from intake: {audience}

Suggested duration: {{
    "lesson_drawing": "60-90 minutes",
    "decision_forcing": "90-120 minutes",
    "caselet": "20-40 minutes"
}}["{case_type}"]

Return JSON with this exact shape:
{{
  "case_summary": "string <= 400 words",
  "learning_objectives": ["objective 1", "objective 2", ...],
  "teaching_strategy": "string covering session format, duration, discussion flow, facilitation notes, common errors",
  "discussion_questions": [
    {{"question": "string", "answer_guide": "string"}}
  ],
  "assessment_methodology": "string",
  "target_audience_and_domains": "string",
  "additional_resources": [
    {{"title": "string", "source": "string", "relevance": "string", "access": "Freely available | Institutional access"}}
  ]
}}

British English. Past tense for descriptive material. Present tense allowed
for instructional statements addressed to the instructor.

=== CONFIRMED CASE DRAFT ===
{full_text[:14000]}

=== SOURCE MATERIAL (for grounding only) ===
{sources_block}
"""
    return _call_json(prompt)


# ---------------------------------------------------------------------------
# Stage 6 — Document export
# ---------------------------------------------------------------------------

def _add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    sizes = {0: 22, 1: 16, 2: 13, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    if color:
        run.font.color.rgb = color
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_body(doc, text):
    """Render multi-paragraph body text. Lines starting with '##' or '###' are
    promoted to sub-headings."""
    if not text:
        return
    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=3)
        elif line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("# "):
            _add_heading(doc, line[2:].strip(), level=1)
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)


def _strip_inline_markers(text):
    """Hide internal markers like [FOOTNOTE NEEDED: term] from the final doc;
    they are surfaced separately as footnote suggestions."""
    if not text:
        return ""
    return re.sub(r"\[FOOTNOTE NEEDED:[^\]]+\]", "", text)


CBC_LOGO_PATH = "attached_assets/cbc_logo_1770210514857.png"
AGK_LOGO_PATH = "attached_assets/agk_logo_1770210514857.png"


def _add_logo_header_docx(doc):
    """Add CBC + AGK logos at the top of a DOCX cover page as a two-cell
    centred table. Silently skips any logo whose file is missing."""
    try:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left_cell, right_cell = table.rows[0].cells
        left_p = left_cell.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        right_p = right_cell.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if os.path.exists(CBC_LOGO_PATH):
            try:
                left_p.add_run().add_picture(CBC_LOGO_PATH, width=Inches(1.4))
            except Exception:
                pass
        if os.path.exists(AGK_LOGO_PATH):
            try:
                right_p.add_run().add_picture(AGK_LOGO_PATH, width=Inches(1.4))
            except Exception:
                pass
        doc.add_paragraph("")
    except Exception:
        pass


def _add_logo_header_pdf(pdf):
    """Place CBC + AGK logos at the top of the current PDF page (left and
    right). Returns the y-coordinate to continue from. Silently skips any
    logo whose file is missing."""
    try:
        page_w = pdf.w
        margin = pdf.l_margin
        logo_w = 32  # mm
        top_y = pdf.get_y()
        if os.path.exists(CBC_LOGO_PATH):
            try:
                pdf.image(CBC_LOGO_PATH, x=margin, y=top_y, w=logo_w)
            except Exception:
                pass
        if os.path.exists(AGK_LOGO_PATH):
            try:
                pdf.image(AGK_LOGO_PATH, x=page_w - margin - logo_w, y=top_y, w=logo_w)
            except Exception:
                pass
        pdf.set_y(top_y + logo_w * 0.6 + 4)
    except Exception:
        pass


def build_case_docx(case_type, intake_metadata, draft_sections, compliance=None):
    """Produce a styled DOCX matching the Lesson-Drawing / Decision-Forcing /
    Caselet template."""
    doc = Document()

    md = intake_metadata or {}
    title = md.get("title") or "Case Study"
    authors = md.get("authors") or ""

    blue = RGBColor(0x1E, 0x3A, 0x8A)

    # Cover page — logos at the very top
    _add_logo_header_docx(doc)
    _add_heading(doc, title, level=0, color=blue)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[Cover image — to be added by author]")
    r.italic = True
    r.font.size = Pt(10)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(
        f"\nCase Number: 2026CBC/AGK/0XX (to be filled by CBC)\n"
        f"Date: {datetime.utcnow().strftime('%d %B %Y')}\n"
    ).font.size = Pt(11)

    if authors:
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ap.add_run(f"Author(s): {authors}")
        run.font.size = Pt(11)
        run.bold = True

    extras = []
    for label, key in [
        ("Initiative", "initiative"),
        ("Region", "region"),
        ("Sector", "sector"),
        ("Time period", "time_period"),
        ("Audience", "audience"),
    ]:
        val = md.get(key)
        if val:
            extras.append(f"{label}: {val}")
    if extras:
        ep = doc.add_paragraph()
        ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep.add_run("\n" + "  ·  ".join(extras)).font.size = Pt(10)

    doc.add_page_break()

    # AGK front matter — Note from Author, Acknowledgements, Disclaimer
    _add_heading(doc, "Note from the Author", level=1, color=blue)
    _add_body(doc, md.get("note_from_author") or "[To be added by the author.]")
    doc.add_paragraph("")

    _add_heading(doc, "Acknowledgements", level=1, color=blue)
    _add_body(doc, md.get("acknowledgements") or "[To be added by the author.]")
    doc.add_paragraph("")

    _add_heading(doc, "Disclaimer", level=1, color=blue)
    _add_body(
        doc,
        "This case has been written for classroom discussion based on "
        "published and primary sources. It is not intended to illustrate "
        "either effective or ineffective handling of a managerial situation.",
    )
    doc.add_page_break()

    # Body sections
    if case_type == "caselet":
        body = (draft_sections or {}).get("caselet", {}).get("text", "")
        _add_body(doc, _strip_inline_markers(body))
    else:
        for sid, sname in CASE_TYPES[case_type]["sections"]:
            sec = (draft_sections or {}).get(sid)
            if not sec or not sec.get("text"):
                continue
            _add_heading(doc, sname, level=1, color=blue)
            _add_body(doc, _strip_inline_markers(sec["text"]))
            doc.add_paragraph("")

    # Footnote suggestions
    if compliance and compliance.get("footnotes"):
        doc.add_page_break()
        _add_heading(doc, "Footnotes (Suggested)", level=1, color=blue)
        for i, fn in enumerate(compliance["footnotes"], 1):
            p = doc.add_paragraph()
            r = p.add_run(f"[{i}] {fn.get('term', '')}: ")
            r.bold = True
            p.add_run(fn.get("footnote_text", ""))

    # References
    if compliance and compliance.get("references"):
        doc.add_page_break()
        _add_heading(doc, "References", level=1, color=blue)
        for ref in compliance["references"]:
            doc.add_paragraph(ref)

    # Study questions for Lesson-Drawing live in the teaching note; for
    # Decision-Forcing they are part of the body. Caselet has them inline too.

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def build_teaching_note_docx(intake_metadata, teaching_note, case_type):
    """Produce a styled DOCX following the AGK Teaching Note template."""
    doc = Document()
    md = intake_metadata or {}
    blue = RGBColor(0x1E, 0x3A, 0x8A)
    case_label = CASE_TYPES.get(case_type, {}).get("label", "Case")

    _add_logo_header_docx(doc)
    _add_heading(doc, "Teaching Note", level=0, color=blue)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(md.get("title") or "Case Study")
    r.font.size = Pt(13)
    r.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"\nCase Number: 2026CBC/AGK/0XX (to be filled later)\n"
        f"Date: {datetime.utcnow().strftime('%d %B %Y')}\n"
        f"Author(s): {md.get('authors', '')}\n"
        f"Case type: {case_label}\n"
    ).font.size = Pt(10)
    doc.add_page_break()

    tn = teaching_note or {}

    if tn.get("case_summary"):
        _add_heading(doc, "Case Summary", level=1, color=blue)
        _add_body(doc, tn["case_summary"])
        doc.add_paragraph("")

    if tn.get("learning_objectives"):
        _add_heading(doc, "Learning Objectives", level=1, color=blue)
        for lo in tn["learning_objectives"]:
            doc.add_paragraph(lo, style="List Bullet")
        doc.add_paragraph("")

    if tn.get("teaching_strategy"):
        _add_heading(doc, "Teaching Strategy", level=1, color=blue)
        _add_body(doc, tn["teaching_strategy"])
        doc.add_paragraph("")

    if tn.get("discussion_questions"):
        _add_heading(doc, "Discussion Questions with Answers", level=1, color=blue)
        for i, dq in enumerate(tn["discussion_questions"], 1):
            p = doc.add_paragraph()
            r = p.add_run(f"Q{i}. {dq.get('question', '')}")
            r.bold = True
            r.font.size = Pt(11)
            _add_body(doc, dq.get("answer_guide", ""))
            doc.add_paragraph("")

    if tn.get("assessment_methodology"):
        _add_heading(doc, "Assessment Methodology", level=1, color=blue)
        _add_body(doc, tn["assessment_methodology"])
        doc.add_paragraph("")

    if tn.get("target_audience_and_domains"):
        _add_heading(doc, "Target Audience and Subject Domains", level=1, color=blue)
        _add_body(doc, tn["target_audience_and_domains"])
        doc.add_paragraph("")

    if tn.get("additional_resources"):
        _add_heading(doc, "Additional Teaching Resources", level=1, color=blue)
        for res in tn["additional_resources"]:
            p = doc.add_paragraph()
            r = p.add_run(res.get("title", ""))
            r.bold = True
            r.font.size = Pt(11)
            doc.add_paragraph(f"Source: {res.get('source', '')}")
            doc.add_paragraph(res.get("relevance", ""))
            access = res.get("access", "")
            if access:
                ap = doc.add_paragraph()
                ar = ap.add_run(f"Access: {access}")
                ar.italic = True
                ar.font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


class _CasePDF(FPDF):
    def header(self):
        if getattr(self, "_skip_header", False):
            return
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(30, 58, 138)
        self.cell(0, 8, sanitize_text_for_pdf(self._doc_title or ""), align="C")
        self.ln(10)
        self.set_text_color(0, 0, 0)

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 8, f"Page {self.page_no()}", align="C")
        self.set_text_color(0, 0, 0)


def build_case_pdf(case_type, intake_metadata, draft_sections, compliance=None):
    """Produce a PDF version of the case for quick preview/download."""
    md = intake_metadata or {}
    title = md.get("title") or "Case Study"

    pdf = _CasePDF()
    pdf._doc_title = title
    pdf._skip_header = True
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # Cover — logos at the top, then title
    _add_logo_header_pdf(pdf)
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(30, 58, 138)
    pdf.multi_cell(0, 10, sanitize_text_for_pdf(title), align="C")
    pdf.ln(6)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "I", 10)
    pdf.multi_cell(0, 6, sanitize_text_for_pdf("[Cover image — to be added by author]"), align="C")
    pdf.ln(8)
    pdf.set_font("Helvetica", "", 11)
    cover_lines = [
        f"Case Number: 2026CBC/AGK/0XX",
        f"Date: {datetime.utcnow().strftime('%d %B %Y')}",
    ]
    if md.get("authors"):
        cover_lines.append(f"Author(s): {md.get('authors')}")
    for label, key in [
        ("Initiative", "initiative"),
        ("Region", "region"),
        ("Sector", "sector"),
        ("Time period", "time_period"),
        ("Audience", "audience"),
    ]:
        val = md.get(key)
        if val:
            cover_lines.append(f"{label}: {val}")
    for line in cover_lines:
        pdf.multi_cell(0, 6, sanitize_text_for_pdf(line), align="C")

    pdf.add_page()
    pdf._skip_header = False

    # AGK front matter — Note from Author, Acknowledgements, Disclaimer
    def _front_matter_block(heading, body):
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(30, 58, 138)
        pdf.multi_cell(0, 7, sanitize_text_for_pdf(heading))
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 6, sanitize_text_for_pdf(body))
        pdf.ln(4)

    _front_matter_block(
        "Note from the Author",
        md.get("note_from_author") or "[To be added by the author.]",
    )
    _front_matter_block(
        "Acknowledgements",
        md.get("acknowledgements") or "[To be added by the author.]",
    )
    _front_matter_block(
        "Disclaimer",
        "This case has been written for classroom discussion based on "
        "published and primary sources. It is not intended to illustrate "
        "either effective or ineffective handling of a managerial situation.",
    )
    pdf.add_page()

    def write_section_body(body):
        pdf.set_font("Helvetica", "", 11)
        for raw in body.split("\n"):
            line = raw.rstrip()
            if not line.strip():
                pdf.ln(3)
                continue
            if line.startswith("### "):
                pdf.set_font("Helvetica", "B", 11)
                pdf.multi_cell(0, 6, sanitize_text_for_pdf(line[4:].strip()))
                pdf.set_font("Helvetica", "", 11)
            elif line.startswith("## "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(30, 58, 138)
                pdf.multi_cell(0, 7, sanitize_text_for_pdf(line[3:].strip()))
                pdf.set_text_color(0, 0, 0)
                pdf.set_font("Helvetica", "", 11)
            else:
                pdf.multi_cell(0, 6, sanitize_text_for_pdf(line))

    if case_type == "caselet":
        body = _strip_inline_markers((draft_sections or {}).get("caselet", {}).get("text", ""))
        write_section_body(body)
    else:
        for sid, sname in CASE_TYPES[case_type]["sections"]:
            sec = (draft_sections or {}).get(sid)
            if not sec or not sec.get("text"):
                continue
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(30, 58, 138)
            pdf.multi_cell(0, 8, sanitize_text_for_pdf(sname))
            pdf.set_text_color(0, 0, 0)
            write_section_body(_strip_inline_markers(sec["text"]))
            pdf.ln(4)

    if compliance and compliance.get("footnotes"):
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(30, 58, 138)
        pdf.multi_cell(0, 8, "Footnotes (Suggested)")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 10)
        for i, fn in enumerate(compliance["footnotes"], 1):
            pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"[{i}] {fn.get('term', '')}: {fn.get('footnote_text', '')}"))

    if compliance and compliance.get("references"):
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(30, 58, 138)
        pdf.multi_cell(0, 8, "References")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 10)
        for ref in compliance["references"]:
            pdf.multi_cell(0, 5, sanitize_text_for_pdf(ref))

    output = pdf.output(dest="S")
    if isinstance(output, str):
        return output.encode("latin-1", errors="replace")
    return bytes(output)


# ---------------------------------------------------------------------------
# Stage 5b — Apply selected compliance fixes back to draft
# ---------------------------------------------------------------------------

def apply_compliance_fixes(case_type, draft_sections, findings, selected_indices):
    """Apply literal find/replace for each selected finding into the matching
    section text. Returns (new_draft_sections, applied_count, skipped).

    A fix is applied only when the original_text appears verbatim in the
    matching section (or anywhere in the document if the section is "Whole
    document"). Skipped fixes are reported with the reason so the UI can
    surface them.
    """
    applied = 0
    skipped = []
    new_sections = {sid: dict(sec or {}) for sid, sec in (draft_sections or {}).items()}

    section_label_to_sid = {}
    if case_type == "caselet":
        section_label_to_sid["Caselet"] = "caselet"
        section_label_to_sid["Whole document"] = "caselet"
    else:
        for sid, sname in CASE_TYPES.get(case_type, {}).get("sections", []):
            section_label_to_sid[sname] = sid

    for idx in selected_indices:
        if idx < 0 or idx >= len(findings):
            continue
        f = findings[idx] or {}
        original = (f.get("original_text") or "").strip()
        suggestion = (f.get("suggestion") or "").strip()
        section_label = (f.get("section") or "").strip()
        if not original or not suggestion:
            skipped.append({"index": idx, "reason": "Original or suggestion is empty"})
            continue

        target_sids = []
        if section_label and section_label in section_label_to_sid:
            target_sids.append(section_label_to_sid[section_label])
        else:
            target_sids = list(new_sections.keys()) if new_sections else (
                ["caselet"] if case_type == "caselet" else
                [s[0] for s in CASE_TYPES.get(case_type, {}).get("sections", [])]
            )

        replaced_anywhere = False
        for sid in target_sids:
            sec = new_sections.get(sid) or {}
            text = sec.get("text", "")
            if not text:
                continue
            if original in text:
                sec["text"] = text.replace(original, suggestion, 1)
                new_sections[sid] = sec
                replaced_anywhere = True
                applied += 1
                break

        if not replaced_anywhere:
            skipped.append({
                "index": idx,
                "reason": "Original phrase not found verbatim — edit manually.",
            })

    return new_sections, applied, skipped


# ---------------------------------------------------------------------------
# Stage 6b — KCM competency mapping (Group 11)
# ---------------------------------------------------------------------------

def _format_kcm_for_prompt():
    """Render the KCM_COMPETENCIES dict as a compact list for the prompt."""
    lines = []
    for category, items in KCM_COMPETENCIES.items():
        lines.append(f"## {category.title()} competencies")
        for key, comp in items.items():
            sub = ", ".join(comp.get("sub_competencies", []) or [])
            lines.append(
                f"- {comp['name']} ({key}): {comp['description']} "
                f"Sub-competencies: {sub}."
            )
    return "\n".join(lines)


def map_kcm_competencies(case_type, draft_sections, intake_metadata):
    """Identify the top 3-4 KCM behavioural and functional competencies woven
    through the case, with a short justification grounded in the draft text.

    Returns ``{"competencies": [{"category": "behavioral|functional", "key": ..., "name": ..., "justification": ...}], "weaving_suggestions": [str]}``
    """
    full_text = _assemble_full_draft(case_type, draft_sections)
    if not full_text.strip():
        return {"competencies": [], "weaving_suggestions": []}

    md = intake_metadata or {}
    declared = md.get("competencies") or ""

    prompt = f"""PROMPT GROUP 11 — Karmayogi Competency Model (KCM) mapping.

Below is the full draft of a case study and the KCM competency catalogue.
Identify the THREE TO FOUR competencies (mix of behavioural and functional)
that the case best illustrates. Pick competencies that are clearly visible in
the protagonist's choices and the events described — do NOT invent
competencies the text does not support.

Author has indicated this competency to align with (may be empty):
"{declared}"

For each competency, give a short justification (max two sentences) that
quotes or paraphrases concrete moments from the draft. Then give 1-3
suggestions for how the author could weave the competency more visibly into
the narrative without changing the underlying facts.

=== KCM CATALOGUE ===
{_format_kcm_for_prompt()}

=== DRAFT ===
{full_text[:18000]}

Return ONLY JSON:
{{
  "competencies": [
    {{
      "category": "behavioral | functional",
      "key": "the dictionary key from the catalogue",
      "name": "the competency display name",
      "justification": "<= 2 sentences grounded in the draft"
    }}
  ],
  "weaving_suggestions": [
    "Suggestion 1...",
    "Suggestion 2..."
  ]
}}
"""
    # Let GeneratorAPIError propagate so the UI can surface the failure
    # explicitly instead of silently showing an empty mapping.
    result = _call_json(prompt)
    return {
        "competencies": result.get("competencies", []) or [],
        "weaving_suggestions": result.get("weaving_suggestions", []) or [],
    }


# ---------------------------------------------------------------------------
# Stage 6 — Plain text export
# ---------------------------------------------------------------------------

def build_case_txt(case_type, intake_metadata, draft_sections):
    """Return the case as a plain-text bytes blob — useful when authors want a
    Unicode-safe copy that bypasses the FPDF latin-1 limitation."""
    md = intake_metadata or {}
    title = md.get("title") or "Case Study"
    lines = [
        title,
        "=" * len(title),
        "",
        f"Case Number: 2026CBC/AGK/0XX (to be filled by CBC)",
        f"Date: {datetime.utcnow().strftime('%d %B %Y')}",
    ]
    if md.get("authors"):
        lines.append(f"Author(s): {md['authors']}")
    if md.get("initiative"):
        lines.append(f"Initiative: {md['initiative']}")
    if md.get("region"):
        lines.append(f"Region: {md['region']}")
    if md.get("sector"):
        lines.append(f"Sector: {md['sector']}")
    lines.append("")
    lines.append("Note from the Author")
    lines.append("-" * 20)
    lines.append(md.get("note_from_author") or "[To be added by the author.]")
    lines.append("")
    lines.append("Acknowledgements")
    lines.append("-" * 16)
    lines.append(md.get("acknowledgements") or "[To be added by the author.]")
    lines.append("")
    lines.append("Disclaimer")
    lines.append("-" * 10)
    lines.append(
        "This case has been written for classroom discussion based on published "
        "and primary sources. It is not intended to illustrate either effective "
        "or ineffective handling of a managerial situation."
    )
    lines.append("")

    if case_type == "caselet":
        body = (draft_sections or {}).get("caselet", {}).get("text", "")
        lines.append(_strip_inline_markers(body))
    else:
        for sid, sname in CASE_TYPES[case_type]["sections"]:
            sec = (draft_sections or {}).get(sid)
            if not sec or not sec.get("text"):
                continue
            lines.append("")
            lines.append(sname)
            lines.append("-" * len(sname))
            lines.append(_strip_inline_markers(sec["text"]))

    return ("\n".join(lines) + "\n").encode("utf-8")
