import os
import json
import PyPDF2
import docx
import requests
import base64
import datetime
from io import BytesIO
from openai import OpenAI
from fpdf import FPDF
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import numpy as np

# Initialize OpenAI client
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
openai_client = OpenAI(api_key=OPENAI_API_KEY)

def normalize_extracted_text(text):
    import re
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'(\w) +- +(\w)', r'\1-\2', text)
    text = re.sub(r'(\w) +-(\w)', r'\1-\2', text)
    text = re.sub(r'(\w)- +(\w)', r'\1-\2', text)
    text = re.sub(r' +([,.:;!?\)])', r'\1', text)
    text = re.sub(r'(\() +', r'\1', text)
    text = re.sub(r'" +', '"', text)
    text = re.sub(r' +"', ' "', text)
    text = re.sub(r'\u2019 +', '\u2019', text)
    text = re.sub(r'\u201c +', '\u201c', text)
    text = re.sub(r' +\u201d', '\u201d', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    lines = text.split('\n')
    cleaned_lines = [line.strip() for line in lines]
    text = '\n'.join(cleaned_lines)
    return text


def extract_text_from_pdf(pdf_file):
    """
    Extract text content from a PDF file.
    
    Args:
        pdf_file: The uploaded PDF file object
        
    Returns:
        str: Extracted text from the PDF
    """
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    return normalize_extracted_text(text)

def extract_text_from_docx(docx_file):
    """
    Extract text content from a DOCX file.
    
    Args:
        docx_file: The uploaded DOCX file object
        
    Returns:
        str: Extracted text from the DOCX
    """
    text = ""
    try:
        # Convert file to BytesIO object
        bytes_content = BytesIO(docx_file.getvalue())
        
        # Open using python-docx
        doc = docx.Document(bytes_content)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    return text

def call_openai_api(prompt, model="gpt-4o", response_format=None, temperature=0.5):
    """
    Call the OpenAI API with the given prompt.
    
    Args:
        prompt (str): The prompt to send to the API
        model (str): The OpenAI model to use
        response_format (str, optional): Format for response (e.g., "json_object")
        temperature (float): Temperature for response generation (default 0.5)
        
    Returns:
        The content of the API response
    """
    try:
        messages = [{"role": "user", "content": prompt}]
        
        kwargs = {}
        if response_format == "json_object":
            kwargs["response_format"] = {"type": "json_object"}
        
        response = openai_client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_tokens=2000,
            **kwargs
        )
        
        # Process the response
        content = response.choices[0].message.content
        
        # If JSON format was requested, parse the result
        if response_format == "json_object":
            try:
                result = json.loads(content)
                
                # Ensure all expected fields exist and have the correct types
                if "score" not in result:
                    result["score"] = 0  # Default to 0 for case study scoring
                elif not isinstance(result["score"], (int, float)):
                    # Convert to number if possible, otherwise use default
                    try:
                        result["score"] = float(result["score"])
                    except ValueError:
                        result["score"] = 0
                
                # Ensure reasoning is a string
                if "reasoning" not in result:
                    result["reasoning"] = "No reasoning provided"
                elif isinstance(result["reasoning"], list):
                    result["reasoning"] = ". ".join([str(item) for item in result["reasoning"]])
                
                # Ensure document_reference is a string
                if "document_reference" not in result:
                    result["document_reference"] = "No document reference provided"
                elif isinstance(result["document_reference"], list):
                    result["document_reference"] = ". ".join([str(item) for item in result["document_reference"]])
                
                return result
            except json.JSONDecodeError:
                return {
                    "score": 0,
                    "reasoning": "Failed to parse response from AI model",
                    "document_reference": "Error in analysis"
                }
        
        return content
    
    except Exception as e:
        return f"Error calling OpenAI API: {str(e)}"


def analyze_writing_quality_chunked(full_text, exclusion_instructions="", chunk_size=3000, overlap=500, max_workers=5):
    from concurrent.futures import ThreadPoolExecutor, as_completed

    text_length = len(full_text)
    if text_length == 0:
        return []

    chunks = []
    start = 0
    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunks.append((start, full_text[start:end]))
        if end >= text_length:
            break
        start += chunk_size - overlap

    total_chunks = len(chunks)

    def process_chunk(chunk_idx, chunk_text):
        prompt = f"""You are an expert English language editor and proofreader specialising in academic and government case studies written in British English.

Analyse the following text chunk (chunk {chunk_idx + 1} of {total_chunks}) for writing issues. Be HIGHLY ACCURATE — only flag genuine errors, not stylistic preferences. Do NOT flag proper nouns, acronyms, abbreviations, or domain-specific terminology.

CHECK FOR THESE SPECIFIC ISSUE TYPES:

1. **Spelling** — Identify words that are misspelt. Flag American English spellings that should be British English (e.g., "organization" → "organisation", "center" → "centre", "analyze" → "analyse", "color" → "colour", "program" → "programme", "defense" → "defence", "favor" → "favour", "realize" → "realise").

2. **Grammar** — Check subject-verb agreement errors, incorrect article usage, missing or extra prepositions, incorrect pronoun references, and misused words.

3. **Tense Consistency** — Case studies should primarily use past tense. Flag any inconsistent shifts to present tense within narrative sections (but allow present tense in general statements, conclusions, or when describing ongoing situations).

4. **Redundancy** — Flag redundant phrases like "in order to" (use "to"), "due to the fact that" (use "because"), "at this point in time" (use "now"), "each and every" (use "each" or "every"), "free gift" (use "gift"), repeated ideas in adjacent sentences.

5. **Sentence Structure** — Flag overly complex sentences (more than 40 words with multiple clauses), sentences in passive voice that would be clearer in active voice, run-on sentences, and suggest simpler alternatives.

{exclusion_instructions}

=== TEXT TO ANALYSE ===
{chunk_text}

Return your findings as a JSON object with this EXACT structure:
{{
    "findings": [
        {{
            "type": "Spelling|Grammar|Tense|Redundancy|Structure",
            "original_text": "the exact problematic text from the document (keep short, max 15 words)",
            "suggestion": "the corrected or improved version",
            "severity": "High|Medium|Low",
            "context": "brief explanation of the issue (one sentence)"
        }}
    ]
}}

Rules:
- Return ONLY genuine issues — do NOT fabricate or invent findings
- If there are no issues in this chunk, return {{"findings": []}}
- Keep original_text short (the specific problematic phrase only, max 15 words)
- severity: High = clear error (spelling, grammar), Medium = tense inconsistency or significant redundancy, Low = style improvement or minor redundancy
- Do NOT flag proper nouns, place names, organisation names, or technical terms
- Do NOT flag direct quotes from other sources
- IMPORTANT: This text was extracted from a PDF. Do NOT flag any spacing-related issues such as extra spaces around hyphens, spaces before punctuation, or irregular whitespace — these are PDF extraction artifacts, NOT errors in the original document. Focus exclusively on genuine content-level spelling, grammar, tense, redundancy, and sentence structure issues.
- Be conservative — when in doubt, do NOT flag it"""

        result = call_openai_api(prompt, response_format="json_object")
        findings = []
        if isinstance(result, dict):
            chunk_findings = result.get("findings", [])
            if isinstance(chunk_findings, list):
                for finding in chunk_findings:
                    if isinstance(finding, dict) and finding.get("original_text"):
                        finding["chunk"] = chunk_idx + 1
                        findings.append(finding)
        return chunk_idx, findings

    all_findings = []
    with ThreadPoolExecutor(max_workers=min(max_workers, total_chunks)) as executor:
        futures = {
            executor.submit(process_chunk, idx, chunk_text): idx
            for idx, (char_offset, chunk_text) in enumerate(chunks)
        }
        for future in as_completed(futures):
            try:
                chunk_idx, findings = future.result()
                all_findings.extend(findings)
            except Exception as e:
                import logging
                logging.warning(f"Writing analysis chunk failed: {str(e)}")

    all_findings.sort(key=lambda f: f.get("chunk", 0))

    seen_texts = set()
    deduplicated = []
    for finding in all_findings:
        orig = finding.get("original_text", "").strip().lower()
        if orig and orig not in seen_texts:
            seen_texts.add(orig)
            deduplicated.append(finding)

    filtered = []
    for finding in deduplicated:
        orig = finding.get("original_text", "").strip()
        suggestion = finding.get("suggestion", "").strip()
        if orig.replace(" ", "").replace("\u00a0", "").lower() == suggestion.replace(" ", "").replace("\u00a0", "").lower() and orig != suggestion:
            continue
        orig_no_hyphspace = orig.replace(" -", "-").replace("- ", "-").replace(" ", "")
        sugg_no_hyphspace = suggestion.replace(" -", "-").replace("- ", "-").replace(" ", "")
        if orig_no_hyphspace.lower() == sugg_no_hyphspace.lower() and orig != suggestion:
            continue
        filtered.append(finding)

    type_order = {"High": 0, "Medium": 1, "Low": 2}
    filtered.sort(key=lambda f: type_order.get(f.get("severity", "Low"), 2))

    return filtered


def create_gauge_chart(score, title, max_value=100):
    """
    Creates a gauge chart for a score
    
    Args:
        score (float): Score value
        title (str): Title for the gauge
        max_value (float): Maximum value for the gauge (default 100 for percentage)
        
    Returns:
        plotly figure object
    """
    # Define color thresholds based on percentage
    threshold_low = max_value * 0.4
    threshold_mid = max_value * 0.7
    
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        title={'text': title},
        domain={'x': [0, 1], 'y': [0, 1]},
        number={'suffix': '%' if max_value == 100 else ''},
        gauge={
            'axis': {'range': [0, max_value], 'tickwidth': 1, 'tickcolor': "darkblue"},
            'bar': {'color': "royalblue"},
            'bgcolor': "white",
            'borderwidth': 2,
            'bordercolor': "gray",
            'steps': [
                {'range': [0, threshold_low], 'color': 'lightcoral'},
                {'range': [threshold_low, threshold_mid], 'color': 'lightyellow'},
                {'range': [threshold_mid, max_value], 'color': 'lightgreen'}
            ]
        }
    ))
    
    fig.update_layout(
        height=300,
        width=400,
        margin=dict(l=30, r=30, b=20, t=40),
        paper_bgcolor="white",
        font={'color': "darkblue", 'family': "Arial"}
    )
    
    return fig
    
def sanitize_text_for_pdf(text):
    """
    Sanitize text to be compatible with FPDF's latin-1 encoding
    
    Args:
        text (str): Text to sanitize
        
    Returns:
        str: Sanitized text
    """
    if not text:
        return ""
    
    # Replace problematic characters
    replacements = {
        # Quotes and apostrophes
        '\u2018': "'",  # Left single quotation mark
        '\u2019': "'",  # Right single quotation mark
        '\u201c': '"',  # Left double quotation mark
        '\u201d': '"',  # Right double quotation mark
        
        # Dashes and hyphens
        '\u2013': '-',  # En dash
        '\u2014': '--', # Em dash
        '\u2015': '--', # Horizontal bar
        '\u2212': '-',  # Minus sign
        
        # Bullets and list markers
        '\u2022': '-',  # Bullet
        '\u2023': '-',  # Triangular bullet
        '\u2043': '-',  # Hyphen bullet
        '\u25e6': '-',  # White bullet
        '\u2219': '-',  # Bullet operator
        
        # Spaces and breaks
        '\u00a0': ' ',  # Non-breaking space
        '\u2000': ' ',  # En quad
        '\u2001': ' ',  # Em quad
        '\u2002': ' ',  # En space
        '\u2003': ' ',  # Em space
        '\u2004': ' ',  # Three-per-em space
        '\u2005': ' ',  # Four-per-em space
        '\u2006': ' ',  # Six-per-em space
        '\u2007': ' ',  # Figure space
        '\u2008': ' ',  # Punctuation space
        '\u2009': ' ',  # Thin space
        '\u200a': ' ',  # Hair space
        
        # Other punctuation
        '\u2026': '...', # Ellipsis
        '\u00b7': '.',   # Middle dot
        
        # Symbols
        '\u2713': 'v',   # Check mark
        '\u2192': '->',  # Rightward arrow
        '\u21e8': '->',  # Rightward white arrow
        '\u25ba': '>'    # Black right-pointing pointer
    }
    
    for original, replacement in replacements.items():
        text = text.replace(original, replacement)
    
    # Remove any other non-latin1 characters
    text = text.encode('latin-1', errors='replace').decode('latin-1')
    
    return text

def generate_report_pdf(filename, document_name, document_summary, assessment_results, assessment_areas, assessment_criteria, recommendations="", weighted_scores=None, competency_mapping=None, tags_data=None, writing_findings=None):
    """
    Generate a PDF report of the case study assessment
    
    Args:
        filename (str): Name to save the PDF as
        document_name (str): Name of the document
        document_summary (str): Summary of the document
        assessment_results (dict): Results of the assessment
        assessment_areas (dict): Assessment areas information
        assessment_criteria (dict): Assessment criteria information
        recommendations (str, optional): Recommendations for improvement
        weighted_scores (dict, optional): Weighted score calculations
        competency_mapping (dict, optional): KCM competency mapping with behavioral and functional competencies
        
    Returns:
        BytesIO: PDF file as BytesIO object
    """
    # Sanitize all text inputs for PDF compatibility
    document_name = sanitize_text_for_pdf(document_name)
    document_summary = sanitize_text_for_pdf(document_summary)
    if recommendations:
        recommendations = sanitize_text_for_pdf(recommendations)
    
    class PDF(FPDF):
        def header(self):
            # Set up the header
            self.set_font('Arial', 'B', 15)
            self.cell(0, 10, 'Case Study Assessment Report', 0, 1, 'C')
            self.ln(5)
        
        def footer(self):
            # Set up the footer
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            date = datetime.datetime.now().strftime("%Y-%m-%d")
            self.cell(0, 10, f'Generated on {date} - Page {self.page_no()}', 0, 0, 'C')
        
        def section_title(self, title):
            # Add a section title
            self.set_font('Arial', 'B', 14)
            self.set_fill_color(230, 230, 250)  # Light purple background
            self.cell(0, 10, title, 0, 1, 'L', True)
            self.ln(5)
        
        def sub_section_title(self, title):
            # Add a subsection title
            self.set_font('Arial', 'B', 12)
            self.cell(0, 10, title, 0, 1, 'L')
            self.ln(2)
        
        def add_score_box(self, label, score, max_score=3):
            # Create a visual score box
            self.set_fill_color(240, 240, 240)
            self.set_font('Arial', 'B', 10)
            self.cell(40, 10, label, 1, 0, 'L')
            
            # Choose color based on score percentage
            score_pct = score / max_score if max_score > 0 else 0
            if score_pct >= 0.75:
                self.set_fill_color(144, 238, 144)  # Light green
            elif score_pct >= 0.5:
                self.set_fill_color(255, 255, 153)  # Light yellow
            else:
                self.set_fill_color(255, 204, 203)  # Light red
                
            self.cell(20, 10, f"{score}/{max_score}", 1, 1, 'C', True)
            self.ln(1)
    
    # Create PDF object
    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Cover page with title and document info
    pdf.set_font('Arial', 'B', 20)
    pdf.cell(0, 20, 'Case Study Assessment Report', 0, 1, 'C')
    pdf.ln(10)
    
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, f"Document: {document_name}", 0, 1, 'C')
    pdf.ln(10)
    
    pdf.set_font('Arial', '', 10)
    pdf.cell(0, 10, f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1, 'C')
    
    # Add Final Composite Score prominently if available
    if weighted_scores:
        pdf.ln(15)
        final_score = weighted_scores.get("final_composite_score", 0)
        
        # Determine grade label
        if final_score >= 85:
            grade_label = "Excellent"
            grade_color = (40, 167, 69)  # Green
        elif final_score >= 70:
            grade_label = "Good"
            grade_color = (59, 130, 246)  # Blue
        elif final_score >= 55:
            grade_label = "Satisfactory"
            grade_color = (255, 193, 7)  # Yellow
        elif final_score >= 40:
            grade_label = "Needs Improvement"
            grade_color = (253, 126, 20)  # Orange
        else:
            grade_label = "Poor"
            grade_color = (220, 53, 69)  # Red
        
        pdf.set_fill_color(*grade_color)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font('Arial', 'B', 18)
        pdf.cell(0, 15, f"Final Composite Score: {final_score:.1f}% - {grade_label}", 0, 1, 'C', True)
        pdf.set_text_color(0, 0, 0)  # Reset to black
    
    # Add logos side by side below the Final Composite Score
    logo_width = 35
    page_width = 210  # A4 width in mm
    total_logo_width = logo_width * 2 + 20  # Two logos with gap
    start_x = (page_width - total_logo_width) / 2
    
    cbc_logo_path = "attached_assets/cbc_logo_1770210514857.png"
    agk_logo_path = "attached_assets/agk_logo_1770210514857.png"
    
    if os.path.exists(cbc_logo_path) and os.path.exists(agk_logo_path):
        pdf.ln(15)
        current_y = pdf.get_y()
        pdf.image(cbc_logo_path, x=start_x, y=current_y, w=logo_width)
        pdf.image(agk_logo_path, x=start_x + logo_width + 20, y=current_y, w=logo_width)
        pdf.ln(40)
    
    # Executive summary page
    pdf.add_page()
    pdf.section_title("Executive Summary")
    pdf.set_font('Arial', '', 11)
    pdf.multi_cell(0, 6, document_summary)
    pdf.ln(5)
    
    if tags_data and isinstance(tags_data, dict):
        sectors = tags_data.get("sectors", [])
        subthemes = tags_data.get("subthemes", {})
        keywords = tags_data.get("keywords", [])
        
        if sectors or keywords:
            pdf.section_title("Sector Classification & Keywords")
            
            if sectors:
                pdf.sub_section_title("Applicable Sectors:")
                for sector in sectors:
                    sector_name = sanitize_text_for_pdf(str(sector))
                    pdf.set_font('Arial', 'B', 10)
                    pdf.set_fill_color(30, 58, 138)
                    pdf.set_text_color(255, 255, 255)
                    pdf.cell(0, 8, f"  {sector_name}", 0, 1, 'L', True)
                    pdf.set_text_color(0, 0, 0)
                    
                    sector_subthemes = subthemes.get(sector, [])
                    if sector_subthemes:
                        pdf.set_font('Arial', '', 9)
                        for sub in sector_subthemes:
                            sub_name = sanitize_text_for_pdf(str(sub))
                            pdf.set_fill_color(240, 240, 250)
                            pdf.cell(10, 6, "", 0, 0)
                            pdf.cell(0, 6, f"- {sub_name}", 0, 1, 'L', True)
                    
                    pdf.ln(2)
                pdf.ln(3)
            
            if keywords:
                pdf.sub_section_title("Keywords:")
                pdf.set_font('Arial', '', 10)
                sanitized_keywords = [sanitize_text_for_pdf(str(k)) for k in keywords]
                keywords_text = ", ".join(sanitized_keywords)
                pdf.multi_cell(0, 6, keywords_text)
                pdf.ln(3)
            
            pdf.ln(5)
    
    # KCM Competency Mapping section
    if competency_mapping and isinstance(competency_mapping, dict):
        pdf.section_title("KCM Competency Mapping")
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, "Based on the case study analysis, the following Karmayogi Competency Model (KCM) competencies have been identified:")
        pdf.ln(3)
        
        # Behavioral Competencies
        behavioral = competency_mapping.get("behavioral_competencies", [])
        if behavioral:
            pdf.set_font('Arial', 'B', 11)
            pdf.set_fill_color(30, 58, 138)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Behavioral Sub-competencies", 0, 1, 'L', True)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(2)
            
            for i, comp in enumerate(behavioral, 1):
                if isinstance(comp, dict):
                    comp_name = sanitize_text_for_pdf(comp.get("name", "Unknown"))
                    parent = sanitize_text_for_pdf(comp.get("parent_competency", ""))
                    display_name = f"{comp_name} ({parent})" if parent else comp_name
                    justification = sanitize_text_for_pdf(comp.get("justification", "No justification provided"))
                    
                    pdf.set_font('Arial', 'B', 10)
                    pdf.set_fill_color(240, 240, 240)
                    pdf.cell(0, 7, f"{i}. {display_name}", 0, 1, 'L', True)
                    
                    pdf.set_font('Arial', '', 9)
                    pdf.multi_cell(0, 5, f"   {justification}")
                    pdf.ln(2)
            pdf.ln(3)
        
        # Functional Competencies
        functional = competency_mapping.get("functional_competencies", [])
        if functional:
            pdf.set_font('Arial', 'B', 11)
            pdf.set_fill_color(30, 58, 138)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Functional Sub-competencies", 0, 1, 'L', True)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(2)
            
            for i, comp in enumerate(functional, 1):
                if isinstance(comp, dict):
                    comp_name = sanitize_text_for_pdf(comp.get("name", "Unknown"))
                    parent = sanitize_text_for_pdf(comp.get("parent_competency", ""))
                    display_name = f"{comp_name} ({parent})" if parent else comp_name
                    justification = sanitize_text_for_pdf(comp.get("justification", "No justification provided"))
                    
                    pdf.set_font('Arial', 'B', 10)
                    pdf.set_fill_color(240, 240, 240)
                    pdf.cell(0, 7, f"{i}. {display_name}", 0, 1, 'L', True)
                    
                    pdf.set_font('Arial', '', 9)
                    pdf.multi_cell(0, 5, f"   {justification}")
                    pdf.ln(2)
            pdf.ln(3)
        
        pdf.ln(5)
    
    # Weighted scoring summary
    if weighted_scores:
        pdf.section_title("Weighted Scoring Summary")
        
        # Create a table for weighted scores
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(70, 8, "Assessment Area", 1, 0, 'L', True)
        pdf.cell(25, 8, "Score", 1, 0, 'C', True)
        pdf.cell(25, 8, "Max", 1, 0, 'C', True)
        pdf.cell(25, 8, "%", 1, 0, 'C', True)
        pdf.cell(25, 8, "Weight", 1, 0, 'C', True)
        pdf.cell(25, 8, "Contrib.", 1, 1, 'C', True)
        
        pdf.set_font('Arial', '', 10)
        for area_id, area_info in assessment_areas.items():
            area_score_data = weighted_scores["area_scores"].get(area_id, {})
            weighted_data = weighted_scores["weighted_scores"].get(area_id, {})
            
            score = area_score_data.get("score", 0)
            max_score = area_score_data.get("max_score", 1)
            percentage = area_score_data.get("percentage", 0)
            weight = weighted_data.get("weight", 0) * 100
            contribution = weighted_data.get("weighted_contribution", 0) * 100
            
            # Shorten long area names
            area_name = area_info["name"]
            if len(area_name) > 35:
                area_name = area_name[:32] + "..."
            
            # Color based on percentage
            if percentage >= 75:
                pdf.set_fill_color(220, 255, 220)
            elif percentage >= 50:
                pdf.set_fill_color(255, 255, 220)
            else:
                pdf.set_fill_color(255, 220, 220)
            
            pdf.cell(70, 8, area_name, 1, 0, 'L')
            pdf.cell(25, 8, f"{score}", 1, 0, 'C')
            pdf.cell(25, 8, f"{max_score}", 1, 0, 'C')
            pdf.cell(25, 8, f"{percentage:.1f}%", 1, 0, 'C', True)
            pdf.cell(25, 8, f"{weight:.0f}%", 1, 0, 'C')
            pdf.cell(25, 8, f"{contribution:.1f}%", 1, 1, 'C')
        
        pdf.ln(5)
        
        # Final score row
        final_score = weighted_scores.get("final_composite_score", 0)
        pdf.set_font('Arial', 'B', 11)
        pdf.cell(145, 10, "Final Composite Score:", 1, 0, 'R')
        
        if final_score >= 75:
            pdf.set_fill_color(144, 238, 144)
        elif final_score >= 50:
            pdf.set_fill_color(255, 255, 153)
        else:
            pdf.set_fill_color(255, 204, 203)
        
        pdf.cell(50, 10, f"{final_score:.1f}%", 1, 1, 'C', True)
        pdf.ln(10)
    
    # Detailed assessment by area
    for area_id, area_info in assessment_areas.items():
        pdf.add_page()
        pdf.section_title(f"Area: {area_info['name']}")
        
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, area_info["description"])
        pdf.ln(3)
        
        # Display area weight and total points
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 8, f"Weight: {area_info.get('weight', 0) * 100:.0f}% | Total Points: {area_info.get('total_points', 0)}", 0, 1)
        pdf.ln(5)
        
        if area_id in assessment_results:
            area_results = assessment_results[area_id]
            
            # Calculate area scores (skip informational criteria)
            area_criteria_defs = assessment_criteria.get(area_id, {})
            total_score = sum(
                result.get("score", 0) for crit_id, result in area_results.items()
                if not area_criteria_defs.get(crit_id, {}).get("informational", False)
            )
            max_possible = area_info.get("total_points", 1)
            percentage = (total_score / max_possible * 100) if max_possible > 0 else 0
            
            # Display area scores in a nice box with a colored background
            if percentage >= 75:
                pdf.set_fill_color(200, 255, 200)  # Light green
            elif percentage >= 50:
                pdf.set_fill_color(255, 255, 200)  # Light yellow
            else:
                pdf.set_fill_color(255, 220, 220)  # Light red
                
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(190, 12, f"Area Score: {total_score}/{max_possible} ({percentage:.1f}%)", 1, 1, 'C', True)
            pdf.ln(5)
            
            # Create a bar chart visualization for the PDF
            pdf.sub_section_title("Criteria Score Visualization")
            
            try:
                # Create a simple visualization using the FPDF directly for reliability
                area_criteria_data = assessment_criteria[area_id]
                
                # Get criteria names and scores
                criteria_list = list(area_results.keys())
                
                # Set up the chart dimensions
                chart_x = 55  # Increased to give more room for labels
                max_bar_width = 100  # Maximum width of bars
                bar_height = 8
                space_between = 5
                header_height = 20
                
                # Calculate total chart height
                chart_height = header_height + (len(criteria_list) * (bar_height + space_between))
                
                # Check if we need a new page
                if pdf.get_y() + chart_height > pdf.page_break_trigger:
                    pdf.add_page()
                
                # Draw chart header
                pdf.set_font('Arial', 'B', 11)
                pdf.cell(0, 10, f"Criteria Scores", 0, 1, 'C') # Centered header
                
                pdf.ln(2)
                
                # Draw each criterion score as a bar (skip informational criteria)
                for crit_id in criteria_list:
                    crit_info = area_criteria_data[crit_id]
                    if crit_info.get("informational", False):
                        continue
                    criterion_name = crit_info["name"]
                    max_score = crit_info.get("max_score", 3)
                    score = area_results[crit_id].get("score", 0)
                    
                    # Draw criterion name
                    bar_y = pdf.get_y()
                    
                    # Position the text before the chart
                    pdf.set_font('Arial', '', 8) # Smaller font for more room
                    
                    # No longer truncating as aggressively, using 45 chars limit
                    if len(criterion_name) > 45:
                        criterion_name = criterion_name[:42] + "..."
                    
                    pdf.set_xy(10, bar_y) # Start from left margin
                    pdf.cell(chart_x - 12, bar_height, criterion_name, 0, 0, 'R')
                    
                    # Draw bar background (gray)
                    pdf.set_fill_color(240, 240, 240)
                    pdf.rect(chart_x, bar_y, max_bar_width, bar_height, 'F')
                    
                    # Choose bar color based on score percentage
                    score_pct = score / max_score if max_score > 0 else 0
                    if score_pct >= 0.75:
                        pdf.set_fill_color(144, 238, 144)  # Light green
                    elif score_pct >= 0.5:
                        pdf.set_fill_color(255, 255, 153)  # Light yellow
                    else:
                        pdf.set_fill_color(255, 204, 203)  # Light red
                    
                    # Draw the score bar
                    bar_width = score_pct * max_bar_width
                    if bar_width > 0:
                        pdf.rect(chart_x, bar_y, bar_width, bar_height, 'F')
                    
                    # Add a black border around the bar
                    pdf.set_draw_color(0, 0, 0)
                    pdf.rect(chart_x, bar_y, max_bar_width, bar_height, 'D')
                    
                    # Add score text
                    pdf.set_font('Arial', 'B', 8)
                    pdf.set_xy(chart_x + max_bar_width + 2, bar_y)
                    pdf.cell(15, bar_height, f"{score}/{max_score}", 0, 1, 'L')
                    
                    # Space for next bar
                    pdf.ln(space_between)
                
                pdf.ln(5)
                
            except Exception as e:
                # If visualization fails, just continue without the chart
                pdf.set_font('Arial', 'I', 10)
                pdf.cell(0, 10, f"Chart visualization not available: {str(e)}", 0, 1)
            
            # Process each criterion with detailed info
            area_criteria = assessment_criteria[area_id]
            
            pdf.sub_section_title("Detailed Criteria Assessment")
            
            for criterion_id, result in area_results.items():
                crit_info = area_criteria[criterion_id]
                criterion_name = crit_info["name"]
                is_informational = crit_info.get("informational", False)

                if is_informational:
                    narrative = result.get("narrative", result.get("reasoning", "No analysis provided"))
                    narrative = sanitize_text_for_pdf(narrative)

                    pdf.set_draw_color(59, 130, 246)
                    pdf.set_fill_color(219, 234, 254)
                    pdf.set_font('Arial', 'B', 11)
                    pdf.cell(0, 8, f"{sanitize_text_for_pdf(criterion_name)} (Informational)", 1, 1, 'L', True)

                    pdf.set_font('Arial', 'I', 9)
                    pdf.set_text_color(59, 130, 246)
                    pdf.cell(0, 6, "This criterion provides informational analysis and is not scored.", 0, 1)
                    pdf.set_text_color(0, 0, 0)

                    pdf.set_font('Arial', 'B', 10)
                    pdf.cell(0, 8, "Analysis:", 0, 1)
                    pdf.set_font('Arial', '', 10)
                    pdf.multi_cell(0, 6, narrative)

                    doc_ref = result.get("document_reference", "")
                    if doc_ref:
                        doc_ref = sanitize_text_for_pdf(doc_ref)
                        pdf.set_font('Arial', 'B', 10)
                        pdf.cell(0, 8, "Document References:", 0, 1)
                        pdf.set_font('Arial', 'I', 9)
                        pdf.multi_cell(0, 6, doc_ref)

                    pdf.set_draw_color(0, 0, 0)
                    pdf.ln(5)
                else:
                    max_score = crit_info.get("max_score", 3)
                    scoring_logic = crit_info.get("scoring_logic", "")
                    score = result.get("score", 0)
                    reasoning = result.get("reasoning", "No reasoning provided")
                    doc_ref = result.get("document_reference", "No references provided")

                    reasoning = sanitize_text_for_pdf(reasoning)
                    doc_ref = sanitize_text_for_pdf(doc_ref)
                    scoring_logic = sanitize_text_for_pdf(scoring_logic)

                    pdf.set_draw_color(100, 100, 100)
                    pdf.set_fill_color(240, 248, 255)
                    pdf.set_font('Arial', 'B', 11)
                    pdf.cell(0, 8, f"{sanitize_text_for_pdf(criterion_name)}", 1, 1, 'L', True)

                    pdf.set_font('Arial', 'I', 9)
                    pdf.cell(0, 6, f"Scoring: {scoring_logic}", 0, 1)

                    score_pct = score / max_score if max_score > 0 else 0
                    if score_pct >= 0.75:
                        pdf.set_fill_color(144, 238, 144)
                    elif score_pct >= 0.5:
                        pdf.set_fill_color(255, 255, 153)
                    else:
                        pdf.set_fill_color(255, 204, 203)

                    pdf.set_font('Arial', 'B', 10)
                    pdf.cell(30, 8, "Score:", 0, 0)
                    pdf.cell(20, 8, f"{score}/{max_score}", 1, 1, 'C', True)

                    pdf.set_font('Arial', 'B', 10)
                    pdf.cell(0, 8, "Reasoning & Evidence:", 0, 1)
                    pdf.set_font('Arial', '', 10)
                    pdf.multi_cell(0, 6, reasoning)

                    pdf.set_font('Arial', 'B', 10)
                    pdf.cell(0, 8, "Document References:", 0, 1)
                    pdf.set_font('Arial', 'I', 9)
                    pdf.multi_cell(0, 6, doc_ref)

                    pdf.ln(5)

            if area_id == "area2" and writing_findings and isinstance(writing_findings, list) and len(writing_findings) > 0:
                pdf.add_page()
                pdf.section_title("Writing Assistant Findings")
                pdf.set_font('Arial', '', 10)
                pdf.multi_cell(0, 5, "Detailed analysis of grammar, spelling, tense consistency, redundancy, and sentence structure across the complete case study.")
                pdf.ln(3)

                type_counts = {}
                severity_counts = {"High": 0, "Medium": 0, "Low": 0}
                for finding in writing_findings:
                    t = finding.get("type", "Other")
                    s = finding.get("severity", "Low")
                    type_counts[t] = type_counts.get(t, 0) + 1
                    if s in severity_counts:
                        severity_counts[s] += 1

                pdf.set_font('Arial', 'B', 10)
                pdf.cell(0, 8, f"Summary: {len(writing_findings)} issues found", 0, 1)
                pdf.set_font('Arial', '', 9)
                summary_parts = []
                for issue_type, count in type_counts.items():
                    summary_parts.append(f"{issue_type}: {count}")
                for sev, count in severity_counts.items():
                    if count > 0:
                        summary_parts.append(f"{sev} severity: {count}")
                pdf.cell(0, 6, " | ".join(summary_parts), 0, 1)
                pdf.ln(5)

                col_widths = [12, 20, 48, 48, 16, 46]
                headers = ["Issue #", "Type", "Original Text", "Suggested Fix", "Severity", "Explanation"]
                line_height = 4

                def draw_table_header():
                    pdf.set_fill_color(30, 58, 138)
                    pdf.set_text_color(255, 255, 255)
                    pdf.set_font('Arial', 'B', 7)
                    for j, header in enumerate(headers):
                        pdf.cell(col_widths[j], 7, header, 1, 0, 'C', True)
                    pdf.ln()
                    pdf.set_text_color(0, 0, 0)

                def calc_lines(text, col_width):
                    if not text:
                        return 1
                    pdf.set_font('Arial', '', 7)
                    char_width = pdf.get_string_width('A')
                    chars_per_line = max(1, int(col_width / char_width))
                    words = text.split()
                    lines = 1
                    current_line_len = 0
                    for word in words:
                        word_len = len(word)
                        if current_line_len == 0:
                            current_line_len = word_len
                        elif current_line_len + 1 + word_len <= chars_per_line:
                            current_line_len += 1 + word_len
                        else:
                            lines += 1
                            current_line_len = word_len
                    return lines

                def draw_wrapped_cell(x, y, w, h, text, align='L', fill=True):
                    pdf.set_xy(x, y)
                    pdf.rect(x, y, w, h, 'DF')
                    pdf.set_xy(x + 1, y + 1)
                    pdf.multi_cell(w - 2, line_height, text, 0, align)

                draw_table_header()

                pdf.set_font('Arial', '', 7)
                for idx, finding in enumerate(writing_findings, 1):
                    sev = finding.get("severity", "Low")
                    orig = sanitize_text_for_pdf(str(finding.get("original_text", "")))
                    sugg = sanitize_text_for_pdf(str(finding.get("suggestion", "")))
                    expl = sanitize_text_for_pdf(str(finding.get("context", "")))
                    f_type = sanitize_text_for_pdf(str(finding.get("type", "Other")))

                    max_lines = max(
                        calc_lines(orig, col_widths[2]),
                        calc_lines(sugg, col_widths[3]),
                        calc_lines(expl, col_widths[5]),
                        1
                    )
                    row_height = max(7, max_lines * line_height + 2)

                    if pdf.get_y() + row_height > pdf.page_break_trigger:
                        pdf.add_page()
                        draw_table_header()

                    if sev == "High":
                        pdf.set_fill_color(248, 215, 218)
                    elif sev == "Medium":
                        pdf.set_fill_color(255, 243, 205)
                    else:
                        pdf.set_fill_color(245, 245, 245)

                    row_y = pdf.get_y()
                    row_x = pdf.get_x()

                    pdf.set_font('Arial', '', 7)
                    x_pos = row_x
                    pdf.rect(x_pos, row_y, col_widths[0], row_height, 'DF')
                    pdf.set_xy(x_pos, row_y + (row_height - line_height) / 2)
                    pdf.cell(col_widths[0], line_height, str(idx), 0, 0, 'C')
                    x_pos += col_widths[0]

                    pdf.rect(x_pos, row_y, col_widths[1], row_height, 'DF')
                    pdf.set_xy(x_pos, row_y + (row_height - line_height) / 2)
                    pdf.cell(col_widths[1], line_height, f_type, 0, 0, 'C')
                    x_pos += col_widths[1]

                    draw_wrapped_cell(x_pos, row_y, col_widths[2], row_height, orig)
                    x_pos += col_widths[2]

                    if sev == "High":
                        pdf.set_fill_color(248, 215, 218)
                    elif sev == "Medium":
                        pdf.set_fill_color(255, 243, 205)
                    else:
                        pdf.set_fill_color(245, 245, 245)

                    draw_wrapped_cell(x_pos, row_y, col_widths[3], row_height, sugg)
                    x_pos += col_widths[3]

                    if sev == "High":
                        pdf.set_fill_color(248, 215, 218)
                    elif sev == "Medium":
                        pdf.set_fill_color(255, 243, 205)
                    else:
                        pdf.set_fill_color(245, 245, 245)

                    pdf.rect(x_pos, row_y, col_widths[4], row_height, 'DF')
                    pdf.set_xy(x_pos, row_y + (row_height - line_height) / 2)
                    pdf.cell(col_widths[4], line_height, sev, 0, 0, 'C')
                    x_pos += col_widths[4]

                    draw_wrapped_cell(x_pos, row_y, col_widths[5], row_height, expl)

                    pdf.set_y(row_y + row_height)

                pdf.ln(5)

        else:
            pdf.set_font('Arial', 'I', 10)
            pdf.cell(0, 10, "No assessment data available for this area", 0, 1)
    
    # Add recommendations section
    if recommendations:
        pdf.add_page()
        pdf.section_title("Recommendations for Improvement")
        pdf.set_font('Arial', '', 11)
        pdf.multi_cell(0, 6, "Based on the assessment results, the following recommendations are provided to improve the case study:")
        pdf.ln(5)
        
        # Format the recommendations text - improved handling
        pdf.set_font('Arial', '', 11)
        
        # Process recommendations with Recommendation X: format
        if "Recommendation" in recommendations:
            # Split by "Recommendation" and process each item
            parts = recommendations.split("Recommendation")
            for part in parts:
                if part.strip():
                    # Clean up the format, preserving sentence integrity
                    cleaned_part = part.strip().replace('\n', ' ')
                    # Remove any stray bullets or formatting that might be causing breaks
                    cleaned_part = cleaned_part.replace('•', '').replace('- ', '')
                    # Sanitize for PDF compatibility
                    cleaned_part = sanitize_text_for_pdf(cleaned_part)
                    
                    # Handle recommendation number formatting
                    if ':' in cleaned_part and cleaned_part.find(':') < 10:
                        num, content = cleaned_part.split(':', 1)
                        
                        # Create a full recommendation text with proper formatting
                        pdf.set_font('Arial', 'B', 11)
                        pdf.write(6, f"Recommendation {num.strip()}: ")
                        
                        # Continue the same line with regular font for content
                        pdf.set_font('Arial', '', 11)
                        pdf.multi_cell(0, 6, content.strip())
                    else:
                        # If no colon format, just print the whole thing
                        pdf.set_font('Arial', 'B', 11)
                        pdf.write(6, "Recommendation: ")
                        pdf.set_font('Arial', '', 11)
                        pdf.multi_cell(0, 6, cleaned_part)
                    
                    pdf.ln(4)
                    
        # Split recommendations by dash list items as a fallback
        elif "-" in recommendations:
            rec_points = recommendations.split("-")
            for point in rec_points:
                if point.strip():
                    # Clean up and replace newlines with spaces for continuous text
                    clean_point = point.strip().replace('\n', ' ')
                    # Sanitize for PDF compatibility
                    clean_point = sanitize_text_for_pdf(clean_point)
                    pdf.set_font('Arial', '', 11)
                    # Use bullet point symbol with proper indentation
                    pdf.cell(5, 6, "-", 0, 0)
                    pdf.multi_cell(0, 6, clean_point)
                    pdf.ln(2)
        # Fallback in case bullet points are used
        elif "•" in recommendations:
            rec_points = recommendations.split("•")
            for point in rec_points:
                if point.strip():
                    # Clean up and replace newlines with spaces for continuous text
                    clean_point = point.strip().replace('\n', ' ')
                    # Sanitize for PDF compatibility
                    clean_point = sanitize_text_for_pdf(clean_point)
                    pdf.set_font('Arial', '', 11)
                    # Use bullet point symbol with proper indentation
                    pdf.cell(5, 6, "-", 0, 0)
                    pdf.multi_cell(0, 6, clean_point)
                    pdf.ln(2)
        else:
            # Just show as regular text if no structured format detected
            # Sanitize for PDF compatibility
            sanitized_recommendations = sanitize_text_for_pdf(recommendations.replace('\n', ' '))
            pdf.multi_cell(0, 6, sanitized_recommendations)
    
    # Save to BytesIO
    pdf_bytes = pdf.output(dest='S')
    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode('latin-1')
    
    pdf_output = BytesIO(pdf_bytes)
    pdf_output.seek(0)
    
    return pdf_output

def get_download_link(pdf_bytes, filename, text):
    """
    Generate a download link for a PDF file
    
    Args:
        pdf_bytes (BytesIO): PDF file as BytesIO object
        filename (str): Filename for download
        text (str): Link text
        
    Returns:
        str: HTML for download link
    """
    b64 = base64.b64encode(pdf_bytes.read()).decode()
    href = f'<a href="data:application/pdf;base64,{b64}" download="{filename}">{text}</a>'
    return href
